# tests/test_workflow.py
import unittest
import os
from datetime import datetime
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.editing.change_model import DocumentChange, ChangeType, ChangeStatus, LocationParagraph, LocationTable, LocationSection
from src.editing.workflow_model import ChangeWorkflow
from src.editing.word_editor import WordEditor # Use the real WordEditor for applier tests
from src.editing.workflow_applier import WorkflowApplier

class TestWorkflowSystem(unittest.TestCase):
    """Unit tests for the ChangeWorkflow model and WorkflowApplier."""

    TEST_DIR = "test_workflow_docs"
    ORIG_DOC_PATH = os.path.join(TEST_DIR, "original_workflow_test.docx")
    MOD_DOC_PATH = os.path.join(TEST_DIR, "modified_workflow_test.docx") # For applier tests

    @classmethod
    def setUpClass(cls):
        """Set up test directory and a sample document for applier tests."""
        if not os.path.exists(cls.TEST_DIR):
            os.makedirs(cls.TEST_DIR)

        # Create a sample document for WorkflowApplier tests
        doc_orig = DocxDocument()
        doc_orig.add_paragraph("Paragraph 0 - Original text.") # Index 0

        p1 = doc_orig.add_paragraph("Paragraph 1 - Text with ") # Index 1
        run1_1 = p1.add_run("bold")
        run1_1.bold = True
        p1.add_run(" and ")
        run1_2 = p1.add_run("italic")
        run1_2.italic = True
        p1.add_run(" formatting.")

        doc_orig.add_paragraph("Paragraph 2 - This paragraph will be deleted.") # Index 2

        doc_orig.add_paragraph("Section Heading", style='Heading 1') # Index 3
        doc_orig.add_paragraph("Content under section heading line 1.") # Index 4
        doc_orig.add_paragraph("Content under section heading line 2.") # Index 5
        doc_orig.add_paragraph("Another Heading", style='Heading 2') # Index 6 - lower level, part of section
        doc_orig.add_paragraph("Content under another heading.") # Index 7
        doc_orig.add_paragraph("Next Section Heading", style='Heading 1') # Index 8

        table = doc_orig.add_table(rows=2, cols=2) # Index 9 (tables are block items like paragraphs)
        table.cell(0, 0).text = "R0C0"
        table.cell(0, 1).text = "R0C1"
        table.cell(1, 0).text = "R1C0"
        table.cell(1, 1).text = "R1C1"

        doc_orig.add_paragraph("Paragraph 10 - Paragraph after table.") # Index 10

        doc_orig.save(cls.ORIG_DOC_PATH)

    @classmethod
    def tearDownClass(cls):
        """Clean up test files and directory."""
        for path in [cls.ORIG_DOC_PATH, cls.MOD_DOC_PATH]:
            if os.path.exists(path):
                os.remove(path)
        if os.path.exists(cls.TEST_DIR):
            if not os.listdir(cls.TEST_DIR): # Only remove if empty
                os.rmdir(cls.TEST_DIR)

    def test_01_workflow_creation_and_add_change(self):
        """Test creating a workflow and adding changes."""
        workflow = ChangeWorkflow(document_id="doc123")
        self.assertEqual(workflow.document_id, "doc123")
        self.assertEqual(len(workflow.proposed_changes), 0)
        self.assertEqual(workflow.status, "active")
        self.assertIsNotNone(workflow.workflow_id)
        self.assertIsNotNone(workflow.created_at)
        self.assertIsNotNone(workflow.updated_at)

        change1 = DocumentChange(
            document_id="doc123",
            change_type=ChangeType.TEXT_UPDATE,
            location=LocationParagraph(paragraph_index=0),
            new_value="new text"
        )
        workflow.add_change(change1)
        self.assertEqual(len(workflow.proposed_changes), 1)
        self.assertEqual(workflow.proposed_changes[0].change_id, change1.change_id)
        self.assertGreaterEqual(workflow.updated_at, workflow.created_at) # Updated timestamp should be later

    def test_02_get_change_by_id(self):
        """Test retrieving a change by its ID."""
        workflow = ChangeWorkflow(document_id="doc123")
        change1 = DocumentChange(document_id="doc123", change_type=ChangeType.TEXT_UPDATE, location=LocationParagraph(0), new_value="new text 1")
        change2 = DocumentChange(document_id="doc123", change_type=ChangeType.PARAGRAPH_DELETE, location=LocationParagraph(1))
        workflow.add_change(change1)
        workflow.add_change(change2)

        retrieved_change = workflow.get_change_by_id(change1.change_id)
        self.assertIsNotNone(retrieved_change)
        self.assertEqual(retrieved_change.change_id, change1.change_id)

        self.assertIsNone(workflow.get_change_by_id("non-existent-id"))

    def test_03_update_change_status(self):
        """Test updating the status of a change."""
        workflow = ChangeWorkflow(document_id="doc123")
        change1 = DocumentChange(document_id="doc123", change_type=ChangeType.TEXT_UPDATE, location=LocationParagraph(0), new_value="new text")
        workflow.add_change(change1)

        initial_updated_at = workflow.updated_at
        self.assertEqual(change1.status, ChangeStatus.PROPOSED)

        success = workflow.update_change_status(change1.change_id, ChangeStatus.ACCEPTED)
        self.assertTrue(success)
        self.assertEqual(change1.status, ChangeStatus.ACCEPTED)
        self.assertGreater(workflow.updated_at, initial_updated_at) # Workflow updated_at should be updated

        success_non_existent = workflow.update_change_status("non-existent-id", ChangeStatus.REJECTED)
        self.assertFalse(success_non_existent)

    def test_04_get_changes_by_status(self):
        """Test retrieving changes filtered by status."""
        workflow = ChangeWorkflow(document_id="doc123")
        change1 = DocumentChange(document_id="doc123", change_type=ChangeType.TEXT_UPDATE, location=LocationParagraph(0), new_value="new text 1")
        change2 = DocumentChange(document_id="doc123", change_type=ChangeType.PARAGRAPH_DELETE, location=LocationParagraph(1))
        change3 = DocumentChange(document_id="doc123", change_type=ChangeType.TEXT_UPDATE, location=LocationParagraph(2), new_value="new text 2")

        workflow.add_change(change1)
        workflow.add_change(change2)
        workflow.add_change(change3)

        self.assertEqual(len(workflow.get_changes_by_status(ChangeStatus.PROPOSED)), 3)
        self.assertEqual(len(workflow.get_changes_by_status(ChangeStatus.ACCEPTED)), 0)
        self.assertEqual(len(workflow.get_changes_by_status(ChangeStatus.REJECTED)), 0)

        workflow.update_change_status(change1.change_id, ChangeStatus.ACCEPTED)
        self.assertEqual(len(workflow.get_changes_by_status(ChangeStatus.PROPOSED)), 2)
        self.assertEqual(len(workflow.get_changes_by_status(ChangeStatus.ACCEPTED)), 1)
        self.assertEqual(workflow.get_accepted_changes()[0].change_id, change1.change_id)

        workflow.update_change_status(change2.change_id, ChangeStatus.REJECTED)
        self.assertEqual(len(workflow.get_changes_by_status(ChangeStatus.PROPOSED)), 1)
        self.assertEqual(len(workflow.get_changes_by_status(ChangeStatus.ACCEPTED)), 1)
        self.assertEqual(len(workflow.get_changes_by_status(ChangeStatus.REJECTED)), 1)
        self.assertEqual(workflow.get_rejected_changes()[0].change_id, change2.change_id)

    def test_05_all_changes_reviewed(self):
        """Test checking if all changes in a workflow have been reviewed."""
        workflow = ChangeWorkflow(document_id="doc123")
        change1 = DocumentChange(document_id="doc123", change_type=ChangeType.TEXT_UPDATE, location=LocationParagraph(0), new_value="new text 1")
        change2 = DocumentChange(document_id="doc123", change_type=ChangeType.PARAGRAPH_DELETE, location=LocationParagraph(1))

        workflow.add_change(change1)
        workflow.add_change(change2)

        self.assertFalse(workflow.all_changes_reviewed())

        workflow.update_change_status(change1.change_id, ChangeStatus.ACCEPTED)
        self.assertFalse(workflow.all_changes_reviewed())

        workflow.update_change_status(change2.change_id, ChangeStatus.REJECTED)
        self.assertTrue(workflow.all_changes_reviewed())

    # --- WorkflowApplier Tests ---

    def test_06_applier_text_update_paragraph(self):
        """Test applying a text update change to a paragraph."""
        # Create a temporary document copy to modify
        temp_doc_path = os.path.join(self.TEST_DIR, "temp_applier_text_para.docx")
        os.system(f"copy {self.ORIG_DOC_PATH} {temp_doc_path}") # Use system copy for simplicity in test setup

        editor = WordEditor(temp_doc_path)
        applier = WorkflowApplier(editor)
        workflow = ChangeWorkflow(document_id="temp_applier_text_para.docx")

        change = DocumentChange(
            document_id="temp_applier_text_para.docx",
            change_type=ChangeType.TEXT_UPDATE,
            location=LocationParagraph(paragraph_index=0),
            old_value="Paragraph 0 - Original text.",
            new_value="Paragraph 0 - Updated by applier.",
            description="Applier test: update paragraph 0"
        )
        workflow.add_change(change)
        workflow.update_change_status(change.change_id, ChangeStatus.ACCEPTED)

        applier.apply_workflow(workflow)

        # Verify the change was applied
        updated_doc = editor.document
        self.assertEqual(updated_doc.paragraphs[0].text, "Paragraph 0 - Updated by applier.")
        self.assertEqual(workflow.get_change_by_id(change.change_id).status, ChangeStatus.APPLIED)

        # Clean up
        if os.path.exists(temp_doc_path):
            os.remove(temp_doc_path)

    def test_07_applier_text_update_run(self):
        """Test applying a text update change to a specific run."""
        temp_doc_path = os.path.join(self.TEST_DIR, "temp_applier_text_run.docx")
        os.system(f"copy {self.ORIG_DOC_PATH} {temp_doc_path}")

        editor = WordEditor(temp_doc_path)
        applier = WorkflowApplier(editor)
        workflow = ChangeWorkflow(document_id="temp_applier_text_run.docx")

        # Change the "bold" run in Paragraph 1 (index 1, run index 1)
        change = DocumentChange(
            document_id="temp_applier_text_run.docx",
            change_type=ChangeType.TEXT_UPDATE,
            location=LocationParagraph(paragraph_index=1, run_index=1),
            old_value="bold",
            new_value="**strong**", # Use markdown-like to check if it's just text replacement
            description="Applier test: update run in paragraph 1"
        )
        workflow.add_change(change)
        workflow.update_change_status(change.change_id, ChangeStatus.ACCEPTED)

        applier.apply_workflow(workflow)

        # Verify the change was applied and formatting preserved
        updated_doc = editor.document
        # The text should be updated
        self.assertEqual(updated_doc.paragraphs[1].runs[1].text, "**strong**")
        # The bold formatting should be preserved
        self.assertTrue(updated_doc.paragraphs[1].runs[1].bold)
        self.assertEqual(workflow.get_change_by_id(change.change_id).status, ChangeStatus.APPLIED)

        # Clean up
        if os.path.exists(temp_doc_path):
            os.remove(temp_doc_path)

    def test_08_applier_paragraph_insert(self):
        """Test applying a paragraph insertion change."""
        temp_doc_path = os.path.join(self.TEST_DIR, "temp_applier_para_insert.docx")
        os.system(f"copy {self.ORIG_DOC_PATH} {temp_doc_path}")

        editor = WordEditor(temp_doc_path)
        applier = WorkflowApplier(editor)
        workflow = ChangeWorkflow(document_id="temp_applier_para_insert.docx")

        # Insert a paragraph before original Paragraph 2 (index 2)
        change = DocumentChange(
            document_id="temp_applier_para_insert.docx",
            change_type=ChangeType.PARAGRAPH_INSERT,
            location=LocationParagraph(paragraph_index=2), # Insert at index 2
            old_value=None,
            new_value="This paragraph was inserted by the applier.",
            description="Applier test: insert paragraph"
        )
        workflow.add_change(change)
        workflow.update_change_status(change.change_id, ChangeStatus.ACCEPTED)

        original_para_count = len(editor.document.paragraphs)

        applier.apply_workflow(workflow)

        # Verify the change was applied
        updated_doc = editor.document
        self.assertEqual(len(updated_doc.paragraphs), original_para_count + 1)
        self.assertEqual(updated_doc.paragraphs[2].text, "This paragraph was inserted by the applier.")
        # Check that the original paragraph at index 2 (Paragraph 2 - This paragraph will be deleted.) is now at index 3
        self.assertEqual(updated_doc.paragraphs[3].text, "Paragraph 2 - This paragraph will be deleted.")
        self.assertEqual(workflow.get_change_by_id(change.change_id).status, ChangeStatus.APPLIED)

        # Clean up
        if os.path.exists(temp_doc_path):
            os.remove(temp_doc_path)

    def test_09_applier_paragraph_delete(self):
        """Test applying a paragraph deletion change."""
        temp_doc_path = os.path.join(self.TEST_DIR, "temp_applier_para_delete.docx")
        os.system(f"copy {self.ORIG_DOC_PATH} {temp_doc_path}")

        editor = WordEditor(temp_doc_path)
        applier = WorkflowApplier(editor)
        workflow = ChangeWorkflow(document_id="temp_applier_para_delete.docx")

        # Delete original Paragraph 2 (index 2)
        change = DocumentChange(
            document_id="temp_applier_para_delete.docx",
            change_type=ChangeType.PARAGRAPH_DELETE,
            location=LocationParagraph(paragraph_index=2), # Delete paragraph at index 2
            old_value="Paragraph 2 - This paragraph will be deleted.",
            new_value=None,
            description="Applier test: delete paragraph 2"
        )
        workflow.add_change(change)
        workflow.update_change_status(change.change_id, ChangeStatus.ACCEPTED)

        original_para_count = len(editor.document.paragraphs)

        applier.apply_workflow(workflow)

        # Verify the change was applied
        updated_doc = editor.document
        self.assertEqual(len(updated_doc.paragraphs), original_para_count - 1)
        # Check that the paragraph at index 2 is now the original paragraph 3
        self.assertEqual(updated_doc.paragraphs[2].text, "Section Heading") # Original index 3
        self.assertEqual(workflow.get_change_by_id(change.change_id).status, ChangeStatus.APPLIED)

        # Clean up
        if os.path.exists(temp_doc_path):
            os.remove(temp_doc_path)

    def test_10_applier_section_replace(self):
        """Test applying a section replacement change."""
        temp_doc_path = os.path.join(self.TEST_DIR, "temp_applier_section_replace.docx")
        os.system(f"copy {self.ORIG_DOC_PATH} {temp_doc_path}")

        editor = WordEditor(temp_doc_path)
        applier = WorkflowApplier(editor)
        workflow = ChangeWorkflow(document_id="temp_applier_section_replace.docx")

        # Replace content under "Section Heading" (original index 3)
        new_content = ["New content line A.", "New content line B."]
        change = DocumentChange(
            document_id="temp_applier_section_replace.docx",
            change_type=ChangeType.SECTION_REPLACE,
            location=LocationSection(heading_text="Section Heading", heading_style_name="Heading 1"),
            old_value=None,
            new_value=new_content,
            description="Applier test: replace section content"
        )
        workflow.add_change(change)
        workflow.update_change_status(change.change_id, ChangeStatus.ACCEPTED)

        applier.apply_workflow(workflow)

        # Verify the change was applied
        updated_doc = editor.document
        # The heading should remain
        self.assertEqual(updated_doc.paragraphs[3].text, "Section Heading")
        # The content after the heading should be the new content
        self.assertEqual(updated_doc.paragraphs[4].text, "New content line A.")
        self.assertEqual(updated_doc.paragraphs[5].text, "New content line B.")
        # The paragraph after the section should be the original "Next Section Heading"
        self.assertEqual(updated_doc.paragraphs[6].text, "Next Section Heading")
        self.assertEqual(workflow.get_change_by_id(change.change_id).status, ChangeStatus.APPLIED)

        # Clean up
        if os.path.exists(temp_doc_path):
            os.remove(temp_doc_path)

    def test_11_applier_table_cell_update(self):
        """Test applying a table cell update change."""
        temp_doc_path = os.path.join(self.TEST_DIR, "temp_applier_table_cell.docx")
        os.system(f"copy {self.ORIG_DOC_PATH} {temp_doc_path}")

        editor = WordEditor(temp_doc_path)
        applier = WorkflowApplier(editor)
        workflow = ChangeWorkflow(document_id="temp_applier_table_cell.docx")

        # Update cell R1C1 (table index 0, row 1, col 1)
        change = DocumentChange(
            document_id="temp_applier_table_cell.docx",
            change_type=ChangeType.TABLE_CELL_UPDATE,
            location=LocationTable(table_index=0, row_index=1, column_index=1),
            old_value="R1C1",
            new_value="Cell Updated by Applier",
            description="Applier test: update table cell"
        )
        workflow.add_change(change)
        workflow.update_change_status(change.change_id, ChangeStatus.ACCEPTED)

        applier.apply_workflow(workflow)

        # Verify the change was applied
        updated_doc = editor.document
        self.assertEqual(updated_doc.tables[0].cell(1, 1).text, "Cell Updated by Applier")
        self.assertEqual(workflow.get_change_by_id(change.change_id).status, ChangeStatus.APPLIED)

        # Clean up
        if os.path.exists(temp_doc_path):
            os.remove(temp_doc_path)

    def test_12_applier_mixed_changes(self):
        """Test applying a mix of accepted changes."""
        temp_doc_path = os.path.join(self.TEST_DIR, "temp_applier_mixed.docx")
        os.system(f"copy {self.ORIG_DOC_PATH} {temp_doc_path}")

        editor = WordEditor(temp_doc_path)
        applier = WorkflowApplier(editor)
        workflow = ChangeWorkflow(document_id="temp_applier_mixed.docx")

        # Define a mix of changes (order matters for this simple applier)
        changes = [
             # PARAGRAPH_DELETE - Delete original Paragraph 2 (index 2) - Applied first
            DocumentChange(
                document_id="temp_applier_mixed.docx",
                change_type=ChangeType.PARAGRAPH_DELETE,
                location=LocationParagraph(paragraph_index=2),
                old_value="Paragraph 2 - This paragraph will be deleted.",
                new_value=None,
                description="Applier test: delete paragraph 2"
            ),
            # PARAGRAPH_INSERT - Insert before original Paragraph 0 (index 0) - Applied second
             DocumentChange(
                document_id="temp_applier_mixed.docx",
                change_type=ChangeType.PARAGRAPH_INSERT,
                location=LocationParagraph(paragraph_index=0), # Insert at index 0
                old_value=None,
                new_value="Inserted paragraph at the beginning.",
                description="Applier test: insert paragraph at start"
            ),
            # TEXT_UPDATE (paragraph) - Update original Paragraph 0 (now at index 1) - Applied third
            DocumentChange(
                document_id="temp_applier_mixed.docx",
                change_type=ChangeType.TEXT_UPDATE,
                location=LocationParagraph(paragraph_index=1), # Original index 0 is now 1
                old_value="Paragraph 0 - Original text.",
                new_value="Paragraph 0 - Updated after insert.",
                description="Applier test: update original paragraph 0"
            ),
             # TABLE_CELL_UPDATE - Update cell R0C0 (table index 0, row 0, col 0) - Applied last
            DocumentChange(
                document_id="temp_applier_mixed.docx",
                change_type=ChangeType.TABLE_CELL_UPDATE,
                location=LocationTable(table_index=0, row_index=0, column_index=0),
                old_value="R0C0",
                new_value="Updated R0C0 by Applier",
                description="Applier test: update table cell R0C0"
            ),
        ]

        # Add and accept all changes
        for change in changes:
            workflow.add_change(change)
            workflow.update_change_status(change.change_id, ChangeStatus.ACCEPTED)

        original_para_count = len(editor.document.paragraphs) # 11

        applier.apply_workflow(workflow)

        # Verify the changes
        updated_doc = editor.document
        
        # Verify paragraph count: original (11) - 1 (deleted) + 1 (inserted) = 11
        self.assertEqual(len(updated_doc.paragraphs), original_para_count)

        # Verify insertion at index 0
        self.assertEqual(updated_doc.paragraphs[0].text, "Inserted paragraph at the beginning.")

        # Verify update of original paragraph 0 (now at index 1)
        self.assertEqual(updated_doc.paragraphs[1].text, "Paragraph 0 - Updated after insert.")

        # Verify deletion of original paragraph 2
        # Original paragraphs: 0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10
        # After delete index 2: 0, 1, 3, 4, 5, 6, 7, 8, 9, 10 (indices shifted)
        # After insert at index 0: INSERTED, 0, 1, 3, 4, 5, 6, 7, 8, 9, 10
        # The paragraph at index 2 should be original paragraph 1
        self.assertEqual(updated_doc.paragraphs[2].text, "Paragraph 1 - Text with bold and italic formatting.")

        # Verify table cell update
        self.assertEqual(updated_doc.tables[0].cell(0, 0).text, "Updated R0C0 by Applier")

        # Verify all accepted changes are marked as APPLIED
        for change in changes:
            self.assertEqual(workflow.get_change_by_id(change.change_id).status, ChangeStatus.APPLIED)

        # Clean up
        if os.path.exists(temp_doc_path):
            os.remove(temp_doc_path)


    def test_13_applier_rejected_changes_not_applied(self):
        """Test that changes with status REJECTED are not applied."""
        temp_doc_path = os.path.join(self.TEST_DIR, "temp_applier_rejected.docx")
        os.system(f"copy {self.ORIG_DOC_PATH} {temp_doc_path}")

        editor = WordEditor(temp_doc_path)
        applier = WorkflowApplier(editor)
        workflow = ChangeWorkflow(document_id="temp_applier_rejected.docx")

        # Create a change and mark it as REJECTED
        change = DocumentChange(
            document_id="temp_applier_rejected.docx",
            change_type=ChangeType.TEXT_UPDATE,
            location=LocationParagraph(paragraph_index=0),
            old_value="Paragraph 0 - Original text.",
            new_value="This text should NOT appear.",
            description="Applier test: rejected change"
        )
        workflow.add_change(change)
        workflow.update_change_status(change.change_id, ChangeStatus.REJECTED)

        original_text = editor.document.paragraphs[0].text

        applier.apply_workflow(workflow)

        # Verify the change was NOT applied
        updated_doc = editor.document
        self.assertEqual(updated_doc.paragraphs[0].text, original_text)
        self.assertEqual(workflow.get_change_by_id(change.change_id).status, ChangeStatus.REJECTED) # Status should remain REJECTED

        # Clean up
        if os.path.exists(temp_doc_path):
            os.remove(temp_doc_path)

    # Add tests for error handling during application if specific error cases are identified.


if __name__ == '__main__':
    unittest.main()