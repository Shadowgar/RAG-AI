# tests/test_change_detector.py
import unittest
import os
from docx import Document as DocxDocument

from src.editing.change_detector import ChangeDetector
from src.editing.change_model import ChangeType, LocationParagraph

class TestChangeDetector(unittest.TestCase):
    """Unit tests for the ChangeDetector class."""

    TEST_DOC_DIR = "test_detector_docs"
    ORIG_DOC_PATH = os.path.join(TEST_DOC_DIR, "original.docx")
    MOD_DOC_PATH_IDENTICAL = os.path.join(TEST_DOC_DIR, "modified_identical.docx")
    MOD_DOC_PATH_TEXT_CHANGE = os.path.join(TEST_DOC_DIR, "modified_text_change.docx")
    MOD_DOC_PATH_PARA_ADDED = os.path.join(TEST_DOC_DIR, "modified_para_added.docx")
    MOD_DOC_PATH_PARA_DELETED = os.path.join(TEST_DOC_DIR, "modified_para_deleted.docx")
    MOD_DOC_PATH_MIXED = os.path.join(TEST_DOC_DIR, "modified_mixed.docx")

    @classmethod
    def setUpClass(cls):
        """Set up test directory and sample documents."""
        if not os.path.exists(cls.TEST_DOC_DIR):
            os.makedirs(cls.TEST_DOC_DIR)

        # Create original document
        doc_orig = DocxDocument()
        doc_orig.add_paragraph("Paragraph 1: Unchanged text.")
        doc_orig.add_paragraph("Paragraph 2: Original text to be changed.")
        doc_orig.add_paragraph("Paragraph 3: This will be deleted in one version.")
        doc_orig.add_paragraph("Paragraph 4: Another unchanged paragraph.")
        doc_orig.save(cls.ORIG_DOC_PATH)

        # Create identical modified document
        doc_identical = DocxDocument()
        doc_identical.add_paragraph("Paragraph 1: Unchanged text.")
        doc_identical.add_paragraph("Paragraph 2: Original text to be changed.")
        doc_identical.add_paragraph("Paragraph 3: This will be deleted in one version.")
        doc_identical.add_paragraph("Paragraph 4: Another unchanged paragraph.")
        doc_identical.save(cls.MOD_DOC_PATH_IDENTICAL)

        # Create document with text change
        doc_text_change = DocxDocument()
        doc_text_change.add_paragraph("Paragraph 1: Unchanged text.")
        doc_text_change.add_paragraph("Paragraph 2: MODIFIED text here.") # Changed
        doc_text_change.add_paragraph("Paragraph 3: This will be deleted in one version.")
        doc_text_change.add_paragraph("Paragraph 4: Another unchanged paragraph.")
        doc_text_change.save(cls.MOD_DOC_PATH_TEXT_CHANGE)

        # Create document with paragraph added at the end
        doc_para_added = DocxDocument()
        doc_para_added.add_paragraph("Paragraph 1: Unchanged text.")
        doc_para_added.add_paragraph("Paragraph 2: Original text to be changed.")
        doc_para_added.add_paragraph("Paragraph 3: This will be deleted in one version.")
        doc_para_added.add_paragraph("Paragraph 4: Another unchanged paragraph.")
        doc_para_added.add_paragraph("Paragraph 5: This is an ADDED paragraph.") # Added
        doc_para_added.save(cls.MOD_DOC_PATH_PARA_ADDED)

        # Create document with paragraph deleted from the end
        doc_para_deleted = DocxDocument()
        doc_para_deleted.add_paragraph("Paragraph 1: Unchanged text.")
        doc_para_deleted.add_paragraph("Paragraph 2: Original text to be changed.")
        doc_para_deleted.add_paragraph("Paragraph 3: This will be deleted in one version.")
        # Paragraph 4 is deleted
        doc_para_deleted.save(cls.MOD_DOC_PATH_PARA_DELETED)
        
        # Create document with mixed changes (for current simple detector)
        doc_mixed = DocxDocument()
        doc_mixed.add_paragraph("Paragraph 1: Unchanged text.")
        doc_mixed.add_paragraph("Paragraph 2: MODIFIED mixed text.") # Changed
        # Paragraph 3 ("This will be deleted...") is omitted
        doc_mixed.add_paragraph("Paragraph 4: Another unchanged paragraph.") # This is now para index 2
        doc_mixed.add_paragraph("Paragraph 5: An added paragraph in mixed.") # Added, now para index 3
        doc_mixed.save(cls.MOD_DOC_PATH_MIXED)


    @classmethod
    def tearDownClass(cls):
        """Clean up test files and directory."""
        for path in [
            cls.ORIG_DOC_PATH, cls.MOD_DOC_PATH_IDENTICAL, 
            cls.MOD_DOC_PATH_TEXT_CHANGE, cls.MOD_DOC_PATH_PARA_ADDED,
            cls.MOD_DOC_PATH_PARA_DELETED, cls.MOD_DOC_PATH_MIXED
        ]:
            if os.path.exists(path):
                os.remove(path)
        if os.path.exists(cls.TEST_DOC_DIR):
            if not os.listdir(cls.TEST_DOC_DIR): # Only remove if empty
                os.rmdir(cls.TEST_DOC_DIR)

    def test_01_identical_documents(self):
        """Test with two identical documents, expecting no changes."""
        detector = ChangeDetector(self.ORIG_DOC_PATH, self.MOD_DOC_PATH_IDENTICAL, "doc_identical")
        changes = detector.detect_paragraph_text_changes()
        self.assertEqual(len(changes), 0)

    def test_02_text_change_only(self):
        """Test with a document that has only text changes in existing paragraphs."""
        detector = ChangeDetector(self.ORIG_DOC_PATH, self.MOD_DOC_PATH_TEXT_CHANGE, "doc_text_change")
        changes = detector.detect_paragraph_text_changes()
        
        self.assertEqual(len(changes), 1)
        change = changes[0]
        self.assertEqual(change.change_type, ChangeType.TEXT_UPDATE)
        self.assertEqual(change.location, LocationParagraph(paragraph_index=1))
        self.assertEqual(change.old_value, "Paragraph 2: Original text to be changed.")
        self.assertEqual(change.new_value, "Paragraph 2: MODIFIED text here.")

    def test_03_paragraph_added_at_end(self):
        """Test with a document that has a paragraph added at the end."""
        detector = ChangeDetector(self.ORIG_DOC_PATH, self.MOD_DOC_PATH_PARA_ADDED, "doc_para_added")
        changes = detector.detect_paragraph_text_changes()
        
        self.assertEqual(len(changes), 1)
        change = changes[0] # The current simple detector finds one "insert"
        self.assertEqual(change.change_type, ChangeType.PARAGRAPH_INSERT)
        self.assertEqual(change.location, LocationParagraph(paragraph_index=4))
        self.assertIsNone(change.old_value)
        self.assertEqual(change.new_value, "Paragraph 5: This is an ADDED paragraph.")

    def test_04_paragraph_deleted_from_end(self):
        """Test with a document that has a paragraph deleted from the end."""
        detector = ChangeDetector(self.ORIG_DOC_PATH, self.MOD_DOC_PATH_PARA_DELETED, "doc_para_deleted")
        changes = detector.detect_paragraph_text_changes()
        
        self.assertEqual(len(changes), 1)
        change = changes[0] # The current simple detector finds one "delete"
        self.assertEqual(change.change_type, ChangeType.PARAGRAPH_DELETE)
        self.assertEqual(change.location, LocationParagraph(paragraph_index=3))
        self.assertEqual(change.old_value, "Paragraph 4: Another unchanged paragraph.")
        self.assertIsNone(change.new_value)

    def test_05_mixed_changes_simple_detection(self):
        """
        Test with mixed changes. Current simple detector will misinterpret due to index shifting.
        Original: P1, P2(orig), P3(del_target), P4(unchanged)
        Modified: P1, P2(mod), P4(now at index 2), P5(added, now at index 3)
        """
        detector = ChangeDetector(self.ORIG_DOC_PATH, self.MOD_DOC_PATH_MIXED, "doc_mixed")
        changes = detector.detect_paragraph_text_changes()
        
        # Expected behavior of the current simple indexed-based diff:
        # 1. P1 vs P1: No change.
        # 2. P2(orig) vs P2(mod): TEXT_UPDATE at index 1.
        #    - old: "Paragraph 2: Original text to be changed."
        #    - new: "Paragraph 2: MODIFIED mixed text."
        # 3. P3(del_target) vs P4(now at index 2 in mod): TEXT_UPDATE at index 2.
        #    - old: "Paragraph 3: This will be deleted in one version."
        #    - new: "Paragraph 4: Another unchanged paragraph."
        # 4. P4(orig, at index 3) vs P5(added, now at index 3 in mod): TEXT_UPDATE at index 3
        #    - old: "Paragraph 4: Another unchanged paragraph."
        #    - new: "Paragraph 5: An added paragraph in mixed."
        # No PARAGRAPH_DELETE or PARAGRAPH_INSERT will be correctly identified for P3 or P5 by this simple logic.
        
        self.assertEqual(len(changes), 3) # Expecting 3 text updates due to simple indexed diff

        # Check change at index 1 (original P2 vs modified P2)
        change1 = next(c for c in changes if c.location == LocationParagraph(paragraph_index=1))
        self.assertEqual(change1.change_type, ChangeType.TEXT_UPDATE)
        self.assertEqual(change1.old_value, "Paragraph 2: Original text to be changed.")
        self.assertEqual(change1.new_value, "Paragraph 2: MODIFIED mixed text.")

        # Check change at index 2 (original P3 vs modified P4)
        change2 = next(c for c in changes if c.location == LocationParagraph(paragraph_index=2))
        self.assertEqual(change2.change_type, ChangeType.TEXT_UPDATE)
        self.assertEqual(change2.old_value, "Paragraph 3: This will be deleted in one version.")
        self.assertEqual(change2.new_value, "Paragraph 4: Another unchanged paragraph.")
        
        # Check change at index 3 (original P4 vs modified P5)
        # The current simple diff will report this as a TEXT_UPDATE on original P4,
        # and then the tail-end logic will report an INSERT for P5 if len_mod > len_orig.
        # Original has 4 paras, Modified has 4 paras. So no tail-end logic fires.
        # The loop `range(common_len)` goes up to `min(4,4) = 4`. So indices 0,1,2,3.
        # Change at index 3:
        change3 = next(c for c in changes if c.location == LocationParagraph(paragraph_index=3))
        self.assertEqual(change3.change_type, ChangeType.TEXT_UPDATE)
        self.assertEqual(change3.old_value, "Paragraph 4: Another unchanged paragraph.")
        self.assertEqual(change3.new_value, "Paragraph 5: An added paragraph in mixed.")

if __name__ == "__main__":
    unittest.main()