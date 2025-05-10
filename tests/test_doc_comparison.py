import pytest
import os
from docx import Document as DocxDocument
from docx.shared import Pt
from src.utils.doc_comparison import compare_documents

# Define paths for dummy test files
TEST_DIR = "tests/temp_comparison_files"
ORIGINAL_DOC_PATH = os.path.join(TEST_DIR, "original.docx")
MODIFIED_DOC_PATH_TEXT_CHANGE = os.path.join(TEST_DIR, "modified_text.docx")
MODIFIED_DOC_PATH_FORMAT_CHANGE = os.path.join(TEST_DIR, "modified_format.docx")
MODIFIED_DOC_PATH_STRUCT_CHANGE = os.path.join(TEST_DIR, "modified_structure.docx")

def create_test_doc(path: str, content_config: list):
    """Helper to create a docx file for testing."""
    doc = DocxDocument()
    for item in content_config:
        p = doc.add_paragraph()
        if isinstance(item, str):
            p.add_run(item)
        elif isinstance(item, list): # list of run configs
            for run_config in item:
                run = p.add_run(run_config.get("text", ""))
                if run_config.get("bold"):
                    run.bold = True
                if run_config.get("italic"):
                    run.italic = True
                if run_config.get("font_name"):
                    run.font.name = run_config.get("font_name")
                if run_config.get("font_size"):
                    run.font.size = Pt(run_config.get("font_size"))
    doc.save(path)

@pytest.fixture(scope="module", autouse=True)
def setup_comparison_test_files():
    """Fixture to create and remove dummy test files for comparison."""
    os.makedirs(TEST_DIR, exist_ok=True)

    # Original Document
    original_content = [
        "This is the first paragraph.",
        [{"text": "This is a "}, {"text": "bold", "bold": True}, {"text": " and "}, {"text": "italic", "italic": True}, {"text": " run."}],
        "Third paragraph with specific font.",
    ]
    create_test_doc(ORIGINAL_DOC_PATH, original_content)
    # Set specific font for the third paragraph's first run in original
    doc_orig_temp = DocxDocument(ORIGINAL_DOC_PATH)
    if len(doc_orig_temp.paragraphs) > 2 and doc_orig_temp.paragraphs[2].runs:
        doc_orig_temp.paragraphs[2].runs[0].font.name = "Arial"
        doc_orig_temp.paragraphs[2].runs[0].font.size = Pt(14)
    doc_orig_temp.save(ORIGINAL_DOC_PATH)


    # Modified Document - Text Change
    modified_content_text = [
        "This is the FIRST paragraph, changed.", # Text change
        [{"text": "This is a "}, {"text": "bold", "bold": True}, {"text": " and "}, {"text": "italic", "italic": True}, {"text": " run."}],
        "Third paragraph with specific font.",
    ]
    create_test_doc(MODIFIED_DOC_PATH_TEXT_CHANGE, modified_content_text)
    doc_mod_text_temp = DocxDocument(MODIFIED_DOC_PATH_TEXT_CHANGE)
    if len(doc_mod_text_temp.paragraphs) > 2 and doc_mod_text_temp.paragraphs[2].runs:
         doc_mod_text_temp.paragraphs[2].runs[0].font.name = "Arial"
         doc_mod_text_temp.paragraphs[2].runs[0].font.size = Pt(14)
    doc_mod_text_temp.save(MODIFIED_DOC_PATH_TEXT_CHANGE)


    # Modified Document - Format Change
    modified_content_format = [
        "This is the first paragraph.",
        [{"text": "This is a "}, {"text": "normal", "bold": False}, {"text": " and "}, {"text": "very italic", "italic": True, "font_name": "Times New Roman"}, {"text": " run."}], # Bold removed, font changed
        "Third paragraph with specific font.",
    ]
    create_test_doc(MODIFIED_DOC_PATH_FORMAT_CHANGE, modified_content_format)
    doc_mod_fmt_temp = DocxDocument(MODIFIED_DOC_PATH_FORMAT_CHANGE)
    if len(doc_mod_fmt_temp.paragraphs) > 2 and doc_mod_fmt_temp.paragraphs[2].runs:
        doc_mod_fmt_temp.paragraphs[2].runs[0].font.name = "Arial" # Keep this same as original for this test
        doc_mod_fmt_temp.paragraphs[2].runs[0].font.size = Pt(12) # Change size
    doc_mod_fmt_temp.save(MODIFIED_DOC_PATH_FORMAT_CHANGE)


    # Modified Document - Structure Change (extra paragraph)
    modified_content_structure = [
        "This is the first paragraph.",
        [{"text": "This is a "}, {"text": "bold", "bold": True}, {"text": " and "}, {"text": "italic", "italic": True}, {"text": " run."}],
        "This is an entirely new third paragraph.",
        "Third paragraph with specific font.", # This becomes the fourth
    ]
    create_test_doc(MODIFIED_DOC_PATH_STRUCT_CHANGE, modified_content_structure)
    doc_mod_struct_temp = DocxDocument(MODIFIED_DOC_PATH_STRUCT_CHANGE)
    # Ensure the "specific font" paragraph (now 4th) has its original font for comparison base
    if len(doc_mod_struct_temp.paragraphs) > 3 and doc_mod_struct_temp.paragraphs[3].runs:
         doc_mod_struct_temp.paragraphs[3].runs[0].font.name = "Arial"
         doc_mod_struct_temp.paragraphs[3].runs[0].font.size = Pt(14)
    doc_mod_struct_temp.save(MODIFIED_DOC_PATH_STRUCT_CHANGE)

    yield

    # Teardown
    if os.path.exists(TEST_DIR):
        import shutil
        shutil.rmtree(TEST_DIR)

def test_compare_identical_documents():
    """Tests comparison of two identical documents."""
    diffs = compare_documents(ORIGINAL_DOC_PATH, ORIGINAL_DOC_PATH)
    assert "No differences found." in diffs[-1] # Last item should be no diffs

def test_compare_text_changes():
    """Tests comparison with text changes."""
    diffs = compare_documents(ORIGINAL_DOC_PATH, MODIFIED_DOC_PATH_TEXT_CHANGE)
    assert any("Text content differs:" in d for d in diffs), "Text difference not detected"
    assert any("- This is the first paragraph." in d for d in diffs)
    assert any("+ This is the FIRST paragraph, changed." in d for d in diffs)

def test_compare_formatting_changes():
    """Tests comparison with formatting changes."""
    diffs = compare_documents(ORIGINAL_DOC_PATH, MODIFIED_DOC_PATH_FORMAT_CHANGE)
    print("Diffs for test_compare_formatting_changes:")
    for d in diffs:
        print(d)
    # Look for specific formatting change reports
    assert any("Bold: True -> None" in d or "Bold: True -> False" in d for d in diffs), "Bold change not detected" # Adjusted for None
    assert any("Font Name: 'None' -> 'Times New Roman'" in d or "Font Name: '' -> 'Times New Roman'" in d for d in diffs), "Font name change not detected"
    # Check for font size change in the third paragraph (which is index 2)
    # The utility reports paragraph numbers 1-based.
    # The specific run is the first run of the third paragraph.
    assert any("Font Size: 14.0pt -> 12.0pt" in d for d in diffs), "Font size change not detected"


def test_compare_structural_changes():
    """Tests comparison with structural changes (e.g., added/removed paragraphs)."""
    diffs = compare_documents(ORIGINAL_DOC_PATH, MODIFIED_DOC_PATH_STRUCT_CHANGE)
    print("Diffs for test_compare_structural_changes:")
    for d in diffs:
        print(d)
    
    # Check that Paragraph 3 header exists and is followed by text difference report
    para3_header_found = False
    text_diff_for_para3_found = False
    current_paragraph_is_3 = False

    for i, d_line in enumerate(diffs):
        if "Paragraph 3:" in d_line:
            current_paragraph_is_3 = True
            para3_header_found = True
        elif "Paragraph " in d_line and "Paragraph 3:" not in d_line: # Moved to another paragraph
            current_paragraph_is_3 = False
        
        if current_paragraph_is_3 and "Text content differs:" in d_line:
            text_diff_for_para3_found = True
            break
            # No need to check further for this specific assertion once text diff for P3 is found

    assert para3_header_found, "Paragraph 3 header not found in diffs"
    assert text_diff_for_para3_found, "Text content difference for Paragraph 3 not found or not correctly associated"
    
    # Check for the extra paragraph indication
    assert any("Paragraph 4: Extra in modified" in d for d in diffs), "Indication of extra paragraph (Paragraph 4) not found"

def test_compare_non_existent_file():
    """Tests comparison when one file does not exist."""
    diffs = compare_documents(ORIGINAL_DOC_PATH, "non_existent_file.docx")
    assert any("Error opening documents:" in d for d in diffs)