import pytest
import os
from unittest.mock import MagicMock, patch
from datetime import datetime

# Assuming components are in src/
from src.rag.indexing import DocumentIndexer
from src.rag.vector_store import VectorStore
from src.rag.embeddings import EmbeddingModel
from src.data.metadata_store import MetadataStore
from src.config import settings

# Define paths for dummy test files
TEST_DIR = "tests/temp_indexing_test_files"
DUMMY_DOCX_PATH = os.path.join(TEST_DIR, "dummy_indexing.docx")
NON_EXISTENT_PATH = os.path.join(TEST_DIR, "non_existent.txt")

@pytest.fixture(scope="module", autouse=True)
def setup_and_teardown_test_files():
    """
    Fixture to create and remove dummy test files for indexer tests.
    """
    os.makedirs(TEST_DIR, exist_ok=True)

    # Create dummy docx
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument()
        doc.add_paragraph("This is the first paragraph for indexing test.")
        doc.add_paragraph("This is the second paragraph.")
        doc.save(DUMMY_DOCX_PATH)
    except ImportError:
        pytest.skip("python-docx not installed, skipping docx tests")
    except Exception as e:
        print(f"Error creating dummy docx: {e}")

    yield # This is where the tests run

    # Teardown: remove test files and directory
    if os.path.exists(TEST_DIR):
        import shutil
        shutil.rmtree(TEST_DIR)

@pytest.fixture
def mock_components():
    """
    Fixture to provide mocked instances of DocumentIndexer dependencies.
    """
    mock_vector_store = MagicMock(spec=VectorStore)
    mock_embedding_model = MagicMock(spec=EmbeddingModel)
    mock_metadata_store = MagicMock(spec=MetadataStore)

    # Configure mocks for expected calls during indexing
    mock_metadata_store.add_document_metadata.return_value = 1 # Simulate returning a document ID
    mock_metadata_store.get_chunks_by_document_id.return_value = [
        {"id": 101, "document_id": 1, "chunk_index": 0, "content": "chunk 1 content", "metadata": {}},
        {"id": 102, "document_id": 1, "chunk_index": 1, "content": "chunk 2 content", "metadata": {}},
    ] # Simulate returning stored chunks with IDs
    mock_embedding_model.generate_embeddings.return_value = [[0.1]*settings.EMBEDDING_BATCH_SIZE, [0.2]*settings.EMBEDDING_BATCH_SIZE] # Dummy embeddings

    return mock_vector_store, mock_embedding_model, mock_metadata_store

# --- Test DocumentIndexer ---

@patch('src.rag.indexing.parse_document')
@patch('src.rag.indexing.chunk_document_elements')
def test_document_indexer_index_document_success(
    mock_chunk_document_elements,
    mock_parse_document,
    mock_components
):
    """
    Tests successful document indexing.
    """
    mock_vector_store, mock_embedding_model, mock_metadata_store = mock_components

    # Configure mocks for this specific test
    mock_parse_document.return_value = [{"content": "element 1"}, {"content": "element 2"}]
    mock_chunk_document_elements.return_value = [
        {"content": "chunk 1 content", "metadata": {}},
        {"content": "chunk 2 content", "metadata": {}},
    ]
    mock_metadata_store.get_chunks_by_document_id.return_value = [ # Ensure this matches chunk_document_elements output count
        {"id": 101, "document_id": 1, "chunk_index": 0, "content": "chunk 1 content", "metadata": {}},
        {"id": 102, "document_id": 1, "chunk_index": 1, "content": "chunk 2 content", "metadata": {}},
    ]
    mock_embedding_model.generate_embeddings.return_value = [[0.1]*settings.EMBEDDING_BATCH_SIZE, [0.2]*settings.EMBEDDING_BATCH_SIZE]


    indexer = DocumentIndexer(mock_vector_store, mock_embedding_model, mock_metadata_store)

    # Ensure dummy file exists before attempting to index
    if not os.path.exists(DUMMY_DOCX_PATH):
         pytest.skip("Dummy docx file not created.")

    indexer.index_document(DUMMY_DOCX_PATH)

    # Assert that the correct methods were called on the mocks
    mock_parse_document.assert_called_once_with(DUMMY_DOCX_PATH)
    mock_chunk_document_elements.assert_called_once() # Check args more specifically if needed
    mock_metadata_store.add_document_metadata.assert_called_once() # Check args more specifically if needed
    mock_metadata_store.add_chunk_metadata.assert_called_once() # Check args more specifically if needed
    mock_metadata_store.get_chunks_by_document_id.assert_called_once() # Check args more specifically if needed
    mock_embedding_model.generate_embeddings.assert_called_once() # Check args more specifically if needed
    mock_vector_store.add_embeddings.assert_called_once() # Check args more specifically if needed

@patch('src.rag.indexing.parse_document')
@patch('src.rag.indexing.chunk_document_elements')
def test_document_indexer_index_document_non_existent_file(
    mock_chunk_document_elements,
    mock_parse_document,
    mock_components
):
    """
    Tests indexing a non-existent file.
    """
    mock_vector_store, mock_embedding_model, mock_metadata_store = mock_components
    indexer = DocumentIndexer(mock_vector_store, mock_embedding_model, mock_metadata_store)

    indexer.index_document(NON_EXISTENT_PATH)

    # Assert that parsing, chunking, metadata, embedding, and vector store methods were NOT called
    mock_parse_document.assert_not_called()
    mock_chunk_document_elements.assert_not_called()
    mock_metadata_store.add_document_metadata.assert_not_called()
    mock_metadata_store.add_chunk_metadata.assert_not_called()
    mock_metadata_store.get_chunks_by_document_id.assert_not_called()
    mock_embedding_model.generate_embeddings.assert_not_called()
    mock_vector_store.add_embeddings.assert_not_called()

@patch('src.rag.indexing.parse_document')
@patch('src.rag.indexing.chunk_document_elements')
def test_document_indexer_index_document_no_elements(
    mock_chunk_document_elements,
    mock_parse_document,
    mock_components
):
    """
    Tests indexing a document where parsing returns no elements.
    """
    mock_vector_store, mock_embedding_model, mock_metadata_store = mock_components
    indexer = DocumentIndexer(mock_vector_store, mock_embedding_model, mock_metadata_store)

    # Configure mocks for this specific test
    mock_parse_document.return_value = [] # Simulate no elements parsed

    # Ensure dummy file exists before attempting to index
    if not os.path.exists(DUMMY_DOCX_PATH):
         pytest.skip("Dummy docx file not created.")

    indexer.index_document(DUMMY_DOCX_PATH)

    # Assert that parsing was called, but subsequent steps were NOT called
    mock_parse_document.assert_called_once_with(DUMMY_DOCX_PATH)
    mock_chunk_document_elements.assert_not_called()
    mock_metadata_store.add_document_metadata.assert_called_once() # Document metadata should still be added
    mock_metadata_store.add_chunk_metadata.assert_not_called()
    mock_metadata_store.get_chunks_by_document_id.assert_not_called()
    mock_embedding_model.generate_embeddings.assert_not_called()
    mock_vector_store.add_embeddings.assert_not_called()

@patch('src.rag.indexing.parse_document')
@patch('src.rag.indexing.chunk_document_elements')
def test_document_indexer_index_document_no_chunks(
    mock_chunk_document_elements,
    mock_parse_document,
    mock_components
):
    """
    Tests indexing a document where chunking returns no chunks.
    """
    mock_vector_store, mock_embedding_model, mock_metadata_store = mock_components
    indexer = DocumentIndexer(mock_vector_store, mock_embedding_model, mock_metadata_store)

    # Configure mocks for this specific test
    mock_parse_document.return_value = [{"content": "element 1"}]
    mock_chunk_document_elements.return_value = [] # Simulate no chunks created

    # Ensure dummy file exists before attempting to index
    if not os.path.exists(DUMMY_DOCX_PATH):
         pytest.skip("Dummy docx file not created.")

    indexer.index_document(DUMMY_DOCX_PATH)

    # Assert that parsing and chunking were called, but subsequent steps were NOT called
    mock_parse_document.assert_called_once_with(DUMMY_DOCX_PATH)
    mock_chunk_document_elements.assert_called_once()
    mock_metadata_store.add_document_metadata.assert_called_once() # Document metadata should still be added
    mock_metadata_store.add_chunk_metadata.assert_called_once() # Chunk metadata should be added (even if empty list)
    mock_metadata_store.get_chunks_by_document_id.assert_called_once() # Should still attempt to get chunks
    mock_embedding_model.generate_embeddings.assert_not_called()
    mock_vector_store.add_embeddings.assert_not_called()

@patch('src.rag.indexing.parse_document')
@patch('src.rag.indexing.chunk_document_elements')
def test_document_indexer_index_document_chunk_mismatch(
    mock_chunk_document_elements,
    mock_parse_document,
    mock_components
):
    """
    Tests indexing where the number of stored chunks doesn't match the number of created chunks.
    """
    mock_vector_store, mock_embedding_model, mock_metadata_store = mock_components
    indexer = DocumentIndexer(mock_vector_store, mock_embedding_model, mock_metadata_store)

    # Configure mocks for this specific test
    mock_parse_document.return_value = [{"content": "element 1"}, {"content": "element 2"}]
    mock_chunk_document_elements.return_value = [
        {"content": "chunk 1 content", "metadata": {}},
        {"content": "chunk 2 content", "metadata": {}},
    ]
    # Simulate a mismatch: get_chunks_by_document_id returns a different number
    mock_metadata_store.get_chunks_by_document_id.return_value = [
        {"id": 101, "document_id": 1, "chunk_index": 0, "content": "chunk 1 content", "metadata": {}},
    ] # Only one chunk returned

    # Ensure dummy file exists before attempting to index
    if not os.path.exists(DUMMY_DOCX_PATH):
         pytest.skip("Dummy docx file not created.")

    indexer.index_document(DUMMY_DOCX_PATH)

    # Assert that parsing, chunking, and metadata calls were made, but embedding and vector store calls were NOT
    mock_parse_document.assert_called_once_with(DUMMY_DOCX_PATH)
    mock_chunk_document_elements.assert_called_once()
    mock_metadata_store.add_document_metadata.assert_called_once()
    mock_metadata_store.add_chunk_metadata.assert_called_once()
    mock_metadata_store.get_chunks_by_document_id.assert_called_once()
    mock_embedding_model.generate_embeddings.assert_not_called() # Embedding should be skipped due to mismatch
    mock_vector_store.add_embeddings.assert_not_called() # Adding to vector store should be skipped