import pytest
import os
import shutil
from unittest.mock import MagicMock, patch
import json

# Assuming modules are in src/rag
from src.rag.embeddings import EmbeddingModel
from src.rag.vector_store import VectorStore
from src.data.metadata_store import MetadataStore
from src.rag.retriever import Retriever
from src.llm.gemini import GeminiClient
from src.llm.prompts import PromptTemplates

# Define paths for temporary test data
TEST_RETRIEVAL_DIR = "tests/temp_retrieval_test_data"
TEST_VECTOR_DB_DIR = os.path.join(TEST_RETRIEVAL_DIR, "vector_db")
TEST_METADATA_DB_PATH = os.path.join(TEST_RETRIEVAL_DIR, "metadata.db")
TEST_DOCUMENTS_DIR = os.path.join(TEST_RETRIEVAL_DIR, "documents")
DUMMY_DOCX_PATH = os.path.join(TEST_DOCUMENTS_DIR, "dummy_retrieval_test.docx")

@pytest.fixture(scope="module", autouse=True)
def setup_and_teardown_retriever_test_data():
    """
    Fixture to create and remove temporary directories and dummy files for Retriever tests.
    """
    os.makedirs(TEST_VECTOR_DB_DIR, exist_ok=True)
    os.makedirs(os.path.dirname(TEST_METADATA_DB_PATH), exist_ok=True)
    os.makedirs(TEST_DOCUMENTS_DIR, exist_ok=True)

    # Create a simple dummy docx file for indexing test
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument()
        doc.add_paragraph("This is a test document for RAG retrieval.")
        doc.add_paragraph("It has a second paragraph about keyword matching.")
        doc.save(DUMMY_DOCX_PATH)
    except ImportError:
        print("python-docx not installed, skipping dummy docx creation.")
    except Exception as e:
        print(f"Error creating dummy docx for RAG test: {e}")

    yield # This is where the tests run

    # Teardown: remove temporary directory
    if os.path.exists(TEST_RETRIEVAL_DIR):
        shutil.rmtree(TEST_RETRIEVAL_DIR)

# Mock dependencies for Retriever test
@pytest.fixture
def mock_embedding_model():
    mock = MagicMock(spec=EmbeddingModel)
    mock.generate_embeddings.side_effect = lambda texts: [[i * 0.1] * 10 for i in range(len(texts))] # Return dummy embeddings
    return mock

@pytest.fixture
def mock_vector_store():
    mock = MagicMock(spec=VectorStore)
    mock.search.return_value = [
        {"external_id": "1", "distance": 0.1},
        {"external_id": "2", "distance": 0.2},
        {"external_id": "3", "distance": 0.3}
    ]
    mock.get_size.return_value = 3
    return mock

@pytest.fixture
def mock_metadata_store():
    mock = MagicMock(spec=MetadataStore)
    mock.get_chunk_by_id.side_effect = lambda chunk_id: {
        "id": chunk_id,
        "content": f"This is chunk content for chunk ID {chunk_id}.",
        "metadata": json.dumps({"page": 1, "source": "test_doc.txt"})
    }
    mock.get_chunks_by_document_id.return_value = [
        {"id": 1, "chunk_index": 0, "content": "chunk 1 content", "metadata": json.dumps({"page": 1})},
        {"id": 2, "chunk_index": 1, "content": "chunk 2 content", "metadata": json.dumps({"page": 1})},
        {"id": 3, "chunk_index": 2, "content": "chunk 3 content", "metadata": json.dumps({"page": 2})},
    ]
    return mock

@pytest.fixture
def mock_gemini_client():
    mock = MagicMock(spec=GeminiClient)
    mock.generate_response.return_value = "This is a generated response from Gemini."
    return mock

def test_retriever_initialization(mock_embedding_model, mock_vector_store, mock_metadata_store, mock_gemini_client):
    retriever = Retriever(mock_embedding_model, mock_vector_store, mock_metadata_store, mock_gemini_client)
    assert retriever.embedding_model == mock_embedding_model
    assert retriever.vector_store == mock_vector_store
    assert retriever.metadata_store == mock_metadata_store
    assert retriever.gemini_client == mock_gemini_client

def test_retriever_retrieve(mock_embedding_model, mock_vector_store, mock_metadata_store, mock_gemini_client):
    retriever = Retriever(mock_embedding_model, mock_vector_store, mock_metadata_store, mock_gemini_client)
    query = "test query"
    results = retriever.retrieve(query, k=2)
    assert isinstance(results, list)
    assert len(results) == 2
    assert "content" in results[0]
    assert "metadata" in results[0]
    assert "response" in results[0]
    assert results[0]["response"] == "This is a generated response from Gemini."

def test_retriever_retrieve_empty_query(mock_embedding_model, mock_vector_store, mock_metadata_store, mock_gemini_client):
    retriever = Retriever(mock_embedding_model, mock_vector_store, mock_metadata_store, mock_gemini_client)
    query = ""
    results = retriever.retrieve(query, k=2)
    assert isinstance(results, list)
    assert len(results) == 0

def test_retriever_retrieve_no_results(mock_embedding_model, mock_metadata_store, mock_gemini_client):
    # Mock vector store to return no results
    mock_vector_store = MagicMock(spec=VectorStore)
    mock_vector_store.search.return_value = []
    mock_vector_store.get_size.return_value = 0

    retriever = Retriever(mock_embedding_model, mock_vector_store, mock_metadata_store, mock_gemini_client)
    query = "test query"
    results = retriever.retrieve(query, k=2)
    assert isinstance(results, list)
    assert len(results) == 0