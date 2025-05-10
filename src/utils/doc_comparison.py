"""
Document Comparison Utility
Provides functions to compare two .docx documents and report differences
in text content and basic formatting.
"""
from typing import List, Dict, Any
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
import difflib

def compare_run_formatting(run1: Run, run2: Run) -> List[str]:
    """Compares basic formatting of two runs."""
    diffs = []
    if run1.bold != run2.bold:
        diffs.append(f"Bold: {run1.bold} -> {run2.bold}")
    if run1.italic != run2.italic:
        diffs.append(f"Italic: {run1.italic} -> {run2.italic}")
    if (run1.font.name or "") != (run2.font.name or ""): # Handle None case
        diffs.append(f"Font Name: '{run1.font.name}' -> '{run2.font.name}'")
    r1_size = run1.font.size.pt if run1.font.size else None
    r2_size = run2.font.size.pt if run2.font.size else None
    if r1_size != r2_size:
        diffs.append(f"Font Size: {r1_size}pt -> {r2_size}pt")
    return diffs

def compare_paragraphs(para1: Paragraph, para2: Paragraph) -> List[str]:
    """Compares two paragraphs run by run."""
    paragraph_diffs = []
    # Simple text comparison for the whole paragraph first
    if para1.text != para2.text:
        # Use difflib for a more granular text diff
        d = difflib.Differ()
        text_diff = list(d.compare(para1.text.splitlines(keepends=True),
                                   para2.text.splitlines(keepends=True)))
        filtered_diff = [line for line in text_diff if line.startswith(('+ ', '- ', '? '))]
        if filtered_diff:
            paragraph_diffs.append("Text content differs:")
            paragraph_diffs.extend([f"  {line.rstrip()}" for line in filtered_diff])

    # Run-by-run comparison for formatting
    runs1 = para1.runs
    runs2 = para2.runs
    max_runs = max(len(runs1), len(runs2))

    for i in range(max_runs):
        run_diffs = []
        if i < len(runs1) and i < len(runs2):
            run1 = runs1[i]
            run2 = runs2[i]
            if run1.text != run2.text:
                run_diffs.append(f"Run text: '{run1.text}' -> '{run2.text}'")
            formatting_diffs = compare_run_formatting(run1, run2)
            run_diffs.extend(formatting_diffs)
        elif i < len(runs1):
            run_diffs.append(f"Extra run in original: '{runs1[i].text}'")
        else: # i < len(runs2)
            run_diffs.append(f"Extra run in modified: '{runs2[i].text}'")

        if run_diffs:
            paragraph_diffs.append(f"  Run {i+1}:")
            paragraph_diffs.extend([f"    - {d}" for d in run_diffs])
    return paragraph_diffs


def compare_documents(original_path: str, modified_path: str) -> List[str]:
    """
    Compares two .docx documents and returns a list of differences.
    Each difference is a string describing the change.
    """
    try:
        doc1 = Document(original_path)
        doc2 = Document(modified_path)
    except Exception as e:
        return [f"Error opening documents: {e}"]

    differences: List[str] = []
    paragraphs1 = doc1.paragraphs
    paragraphs2 = doc2.paragraphs
    max_paras = max(len(paragraphs1), len(paragraphs2))

    differences.append(f"Comparing '{original_path}' and '{modified_path}'")
    differences.append("-" * 30)

    for i in range(max_paras):
        para_diff_details = []
        if i < len(paragraphs1) and i < len(paragraphs2):
            para_diff_details = compare_paragraphs(paragraphs1[i], paragraphs2[i])
            if para_diff_details:
                differences.append(f"Paragraph {i+1}:")
                differences.extend(para_diff_details)
        elif i < len(paragraphs1):
            differences.append(f"Paragraph {i+1}: Extra in original: \"{paragraphs1[i].text[:100]}...\"")
        else: # i < len(paragraphs2)
            differences.append(f"Paragraph {i+1}: Extra in modified: \"{paragraphs2[i].text[:100]}...\"")

    if not any(d for d in differences if not d.startswith("Comparing") and not d.startswith("---")):
        differences.append("No differences found.")

    return differences

if __name__ == '__main__':
    # Basic test usage (requires dummy files to be created)
    # You would create dummy1.docx and dummy2.docx for this to run
    # For example:
    # doc1 = Document()
    # doc1.add_paragraph("This is the first paragraph.")
    # p2 = doc1.add_paragraph("This is ")
    # r = p2.add_run("bold")
    # r.bold = True
    # p2.add_run(" text.")
    # doc1.save("dummy1.docx")

    # doc2 = Document()
    # doc2.add_paragraph("This is the first paragraph, slightly changed.")
    # p3 = doc2.add_paragraph("This is ")
    # r2 = p3.add_run("italic") # Changed from bold to italic
    # r2.italic = True
    # p3.add_run(" text.")
    # doc2.save("dummy2.docx")

    # print("\nComparing dummy1.docx and dummy1.docx (should be no differences):")
    # for diff in compare_documents("dummy1.docx", "dummy1.docx"):
    #     print(diff)
    
    # print("\nComparing dummy1.docx and dummy2.docx:")
    # for diff in compare_documents("dummy1.docx", "dummy2.docx"):
    #     print(diff)
    pass # Keep the if __name__ block for potential local testing