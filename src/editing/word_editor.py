# src/editing/word_editor.py
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from typing import List, Optional
from docx.text.run import Run
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
from ..editing.change_model import DocumentChange, ChangeType, LocationParagraph, LocationTable, LocationSection

class WordEditor:
    """
    Provides utilities for editing Word documents (.docx) while attempting to preserve formatting.
    """

    def __init__(self, doc_path: Optional[str] = None):
        """
        Initializes the WordEditor.

        Args:
            doc_path (str, optional): Path to the Word document. Defaults to None.
                                      If None, a new document is created.
        """
        if doc_path:
            self.document = Document(doc_path)
        else:
            self.document = Document()

    def save_document(self, save_path: str):
        """
        Saves the document to the specified path.

        Args:
            save_path (str): The path where the document will be saved.
        """
        self.document.save(save_path)

    # Placeholder for section-based editing functions
    def replace_text_after_heading(self, heading_text: str, new_content_paragraphs: list[str], heading_style_name: Optional[str] = None):
        """
        Replaces all paragraphs following a specific heading until the next heading
        of the same or higher level, or the end of the document.

        Args:
            heading_text (str): The text of the heading to find.
            new_content_paragraphs (list[str]): A list of strings, where each string
                                                is a new paragraph to insert.
            heading_style_name (str, optional): The style name of the heading
                                                (e.g., 'Heading 1'). If None, matches by text only.
        Returns:
            bool: True if the heading was found and content replaced, False otherwise.
        """
        start_index = -1
        heading_level = -1

        for i, para in enumerate(self.document.paragraphs):
            is_heading_match = False
            if heading_style_name:
                if para.style and para.style.name and para.style.name.startswith(heading_style_name) and heading_text in para.text:
                    is_heading_match = True
                    try:
                        if para.style.name: # Additional check for safety
                            level_str = para.style.name.split(" ")[-1]
                            if level_str.isdigit():
                                heading_level = int(level_str)
                    except:
                        pass
            elif heading_text == para.text.strip():
                is_heading_match = True

            if is_heading_match:
                start_index = i
                break

        if start_index == -1:
            print(f"Heading '{heading_text}' not found.")
            return False

        end_index = len(self.document.paragraphs)
        for i in range(start_index + 1, len(self.document.paragraphs)):
            para = self.document.paragraphs[i]
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                current_level = -1
                try:
                    if para.style.name: # Additional check for safety
                        level_str = para.style.name.split(" ")[-1]
                        if level_str.isdigit():
                            current_level = int(level_str)
                except:
                    pass

                if heading_style_name and heading_level != -1 and current_level != -1:
                    if current_level <= heading_level:
                        end_index = i
                        break
                elif para.style and para.style.name and para.style.name.startswith("Heading"):
                    end_index = i
                    break

        for i in range(end_index - 1, start_index, -1):
            p_to_delete = self.document.paragraphs[i]._element
            p_to_delete.getparent().remove(p_to_delete)

        if new_content_paragraphs:
            if start_index + 1 < len(self.document.paragraphs):
                anchor_for_insertion = self.document.paragraphs[start_index + 1]
                for new_para_text in new_content_paragraphs:
                    new_p = anchor_for_insertion.insert_paragraph_before(new_para_text)
            else:
                for new_para_text in new_content_paragraphs:
                    new_p = self.document.add_paragraph(new_para_text)

        return True


    def insert_paragraph_after(self, paragraph_index: int, text: str, style: Optional[str] = None):
        """
        Inserts a new paragraph with the given text and optional style
        after the paragraph at the specified index.

        Args:
            paragraph_index (int): The index of the paragraph after which to insert.
            text (str): The text for the new paragraph.
            style (str, optional): The style for the new paragraph. Defaults to None.
        Returns:
            Paragraph: The newly inserted paragraph object, or None if index is invalid.
        """
        if not (0 <= paragraph_index < len(self.document.paragraphs)):
            print(f"Warning: Paragraph index {paragraph_index} for insertion is out of bounds.")
            return None
        
        # To insert after paragraph_index, we need to get its element
        # and then use XML manipulation to insert a new <w:p> after it.
        # A simpler way with python-docx is to add to the end if it's the last,
        # or insert_paragraph_before the *next* paragraph.

        if paragraph_index == len(self.document.paragraphs) - 1:
            # If it's the last paragraph, just add a new one
            new_para = self.document.add_paragraph(text, style)
        else:
            # Insert before the paragraph that is currently at paragraph_index + 1
            anchor_paragraph = self.document.paragraphs[paragraph_index + 1]
            new_para = anchor_paragraph.insert_paragraph_before(text, style)
        return new_para

    def delete_paragraph(self, paragraph_index: int) -> bool:
        """
        Deletes the paragraph at the specified index.

        Args:
            paragraph_index (int): The index of the paragraph to delete.
        Returns:
            bool: True if deletion was successful, False otherwise.
        """
        if not (0 <= paragraph_index < len(self.document.paragraphs)):
            print(f"Warning: Paragraph index {paragraph_index} for deletion is out of bounds.")
            return False
        
        p_to_delete = self.document.paragraphs[paragraph_index]
        element = p_to_delete._element
        element.getparent().remove(element)
        return True

    # Table Manipulation Utilities (Placeholders)
    def get_table_cell_text(self, table_index: int, row: int, col: int) -> str:
        """
        Gets the text from a specific cell in a table.
        (Placeholder for implementation)
        """
        # Reason: Placeholder for future implementation.
        if not (0 <= table_index < len(self.document.tables)):
            print(f"Warning: Table index {table_index} out of bounds.")
            return ""
        
        table = self.document.tables[table_index]
        
        if not (0 <= row < len(table.rows)):
            print(f"Warning: Row index {row} out of bounds for table {table_index} with {len(table.rows)} rows.")
            return ""
        
        # A row might not have all columns physically present if they are empty,
        # but table.columns gives the defined number of columns.
        # table.cell(row, col) can create cells if they don't exist up to table.columns.
        if not (0 <= col < len(table.columns)):
            print(f"Warning: Column index {col} out of bounds for table {table_index} with {len(table.columns)} columns.")
            return ""
            
        try:
            return table.cell(row, col).text
        except IndexError:
            # This might happen for jagged tables or other rare conditions not caught by above checks.
            print(f"Warning: Cell ({row}, {col}) caused IndexError for table {table_index}.")
            return ""

    def update_table_cell_text(self, table_index: int, row: int, col: int, text: str) -> bool:
        """
        Updates the text in a specific cell in a table.
        Ensures cell exists before updating.
        """
        if not (0 <= table_index < len(self.document.tables)):
            print(f"Warning: Table index {table_index} out of bounds for update.")
            return False
            
        table = self.document.tables[table_index]

        if not (0 <= row < len(table.rows)):
            print(f"Warning: Row index {row} out of bounds for table {table_index} with {len(table.rows)} rows during update.")
            return False

        if not (0 <= col < len(table.columns)):
            print(f"Warning: Column index {col} out of bounds for table {table_index} with {len(table.columns)} columns during update.")
            return False
            
        try:
            # Accessing cell can create it if it doesn't exist but is within table dimensions
            cell_to_update = table.cell(row, col)
            cell_to_update.text = text
            return True
        except IndexError:
            # Should ideally be caught by above checks, but as a safeguard.
            print(f"Warning: Cell ({row}, {col}) caused IndexError for table {table_index} during update.")
            return False

    def add_row_to_table(self, table_index: int):
        """
        Adds a new row to the specified table. Returns the new row object.
        (Placeholder for implementation)
        """
        # Reason: Placeholder for future implementation.
        if 0 <= table_index < len(self.document.tables):
            table = self.document.tables[table_index]
            return table.add_row()
        return None

    # List Manipulation Utilities (Placeholders - basic, relies on paragraph styles)
    def get_list_item_text(self, paragraph_index: int) -> str:
        """
        Gets the text of a paragraph assumed to be a list item.
        (Placeholder - relies on paragraph styles for true list context)
        """
        # Reason: Placeholder for future implementation.
        # For now, this is identical to get_paragraph_text.
        # True list handling would involve checking paragraph.style and list properties.
        return self.get_paragraph_text(paragraph_index)

    def update_list_item_text(self, paragraph_index: int, text: str, preserve_style: bool = True):
        """
        Updates the text of a paragraph assumed to be a list item.
        (Placeholder - relies on paragraph styles for true list context)
        """
        # Reason: Placeholder for future implementation.
        # For now, this is identical to update_paragraph_text.
        self.update_paragraph_text(paragraph_index, text, preserve_style)
        # True list handling might involve ensuring list style is maintained.


    # Placeholder for paragraph manipulation utilities
    def get_paragraph_text(self, paragraph_index: int) -> str:
        """
        Gets the text of a specific paragraph.
        (Implementation to be developed)
        """
        # Reason: Placeholder for future implementation.
        if 0 <= paragraph_index < len(self.document.paragraphs):
            return self.document.paragraphs[paragraph_index].text
        return "" # Return empty string or raise error for invalid index

    def update_paragraph_text(self, paragraph_index: int, new_text: str, preserve_style: bool = True):
        """
        Updates the text of a specific paragraph, optionally preserving its style.

        Args:
            paragraph_index (int): The index of the paragraph to update.
            new_text (str): The new text for the paragraph.
            preserve_style (bool): If True, attempts to preserve the style of the
                                   first run of the original paragraph. Defaults to True.
        """
        if not (0 <= paragraph_index < len(self.document.paragraphs)):
            # Or raise ValueError("Paragraph index out of bounds")
            print(f"Warning: Paragraph index {paragraph_index} is out of bounds.")
            return

        paragraph = self.document.paragraphs[paragraph_index]

        if preserve_style:
            original_runs = list(paragraph.runs) # Make a copy of run objects
            
            # Extract formatting from the first original run BEFORE clearing
            # This is crucial because run properties might become unreliable after element removal
            run_style_to_apply = None
            font_to_apply = {}
            direct_formatting_to_apply = {}

            if original_runs:
                first_original_run = original_runs[0]
                direct_formatting_to_apply['bold'] = first_original_run.bold
                direct_formatting_to_apply['italic'] = first_original_run.italic
                direct_formatting_to_apply['underline'] = first_original_run.underline
                
                font_to_apply['name'] = first_original_run.font.name
                font_to_apply['size'] = first_original_run.font.size
                font_to_apply['color_rgb'] = first_original_run.font.color.rgb if first_original_run.font.color else None
                font_to_apply['highlight_color'] = first_original_run.font.highlight_color
                font_to_apply['strike'] = first_original_run.font.strike
                font_to_apply['subscript'] = first_original_run.font.subscript
                font_to_apply['superscript'] = first_original_run.font.superscript
                font_to_apply['all_caps'] = first_original_run.font.all_caps
                font_to_apply['small_caps'] = first_original_run.font.small_caps

                if first_original_run.style and first_original_run.style.name:
                    run_style_to_apply = first_original_run.style.name
            else: # No original runs, nothing to copy
                pass


            # Clear existing runs from the live paragraph object
            # Iterate backwards to avoid issues with list modification during iteration
            # This ensures that the paragraph object is clean before adding the new run.
            for r in reversed(list(paragraph.runs)): # Iterate over a copy
                p_elem = r._element.getparent()
                p_elem.remove(r._element)
            # At this point, paragraph.runs should be effectively empty from an XML perspective.
            
            # Add the new text, which creates a new single run
            new_run = paragraph.add_run(new_text)

            # Apply the extracted formatting only if there was something to apply
            if original_runs: # Check if we actually extracted anything
                
                # Apply font attributes first
                if font_to_apply.get('name'): new_run.font.name = font_to_apply['name']
                if font_to_apply.get('size'): new_run.font.size = font_to_apply['size']
                if font_to_apply.get('color_rgb'): new_run.font.color.rgb = font_to_apply['color_rgb']
                if font_to_apply.get('highlight_color') is not None: new_run.font.highlight_color = font_to_apply['highlight_color']
                if font_to_apply.get('strike') is not None: new_run.font.strike = font_to_apply['strike']
                if font_to_apply.get('subscript') is not None: new_run.font.subscript = font_to_apply['subscript']
                if font_to_apply.get('superscript') is not None: new_run.font.superscript = font_to_apply['superscript']
                if font_to_apply.get('all_caps') is not None: new_run.font.all_caps = font_to_apply['all_caps']
                if font_to_apply.get('small_caps') is not None: new_run.font.small_caps = font_to_apply['small_caps']

                # Then, apply character style (if any)
                if run_style_to_apply:
                    try:
                        new_run.style = run_style_to_apply
                    except KeyError:
                        print(f"Warning: Character style '{run_style_to_apply}' not found in document. Skipping style copy for new run.")

                # Finally, apply direct formatting attributes like bold/italic, to ensure they override styles
                # print(f"DEBUG: Applying direct formatting. Bold from source: {direct_formatting_to_apply.get('bold')}")
                
                bold_val = direct_formatting_to_apply.get('bold')
                # TODO: Investigate why new_run.bold = True does not always result in new_run.bold being True
                # when the paragraph has certain styles (e.g., 'ListNumber').
                # It works for default styled paragraphs. This suggests an interaction
                # with how python-docx handles overriding paragraph style with direct run formatting
                # in this specific update sequence (clearing runs, adding new, applying cached format).
                if bold_val is True:
                    new_run.bold = True
                elif bold_val is False:
                    new_run.bold = False
                # If bold_val is None, new_run inherits, so no explicit assignment needed for None.
                    
                italic_val = direct_formatting_to_apply.get('italic')
                if italic_val is True:
                    new_run.italic = True
                elif italic_val is False:
                    new_run.italic = False

                underline_val = direct_formatting_to_apply.get('underline')
                if underline_val is True:
                    new_run.underline = True
                elif underline_val is False:
                    new_run.underline = False
                # Note: python-docx run.underline can also take WD_UNDERLINE enum values
                # For simplicity, sticking to True/False/None for now from direct copy.
            
            # Paragraph's own style is generally preserved unless explicitly changed.

        else:
            # Replace text without preserving style (will use default or paragraph's current style)
            paragraph.text = new_text
            # Ensure paragraph.text clears previous runs and their specific formatting.
            # python-docx handles this by replacing the content of the paragraph.

    def _update_run_text(self, paragraph: Paragraph, run_index: int, new_text: str):
        """
        Updates the text of a specific run within a paragraph while preserving its formatting.

        Args:
            paragraph (Paragraph): The paragraph containing the run.
            run_index (int): The index of the run to update.
            new_text (str): The new text for the run.
        """
        if not (0 <= run_index < len(paragraph.runs)):
            print(f"Warning: Run index {run_index} is out of bounds for paragraph.")
            return

        run = paragraph.runs[run_index]
        original_formatting = {
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'font_name': run.font.name,
            'font_size': run.font.size,
            'color_rgb': run.font.color.rgb if run.font.color else None,
            'highlight_color': run.font.highlight_color,
            'strike': run.font.strike,
            'subscript': run.font.subscript,
            'superscript': run.font.superscript,
            'all_caps': run.font.all_caps,
            'small_caps': run.font.small_caps,
            'style_name': run.style.name if run.style else None
        }

        # Clear the existing run's content
        run.text = ""

        # Add the new text to the run
        run.add_text(new_text)

        # Reapply the original formatting
        run.bold = original_formatting['bold']
        run.italic = original_formatting['italic']
        run.underline = original_formatting['underline']
        run.font.name = original_formatting['font_name']
        if original_formatting['font_size'] is not None:
            run.font.size = original_formatting['font_size']
        if original_formatting['color_rgb'] is not None:
            run.font.color.rgb = original_formatting['color_rgb']
        if original_formatting['highlight_color'] is not None:
             run.font.highlight_color = original_formatting['highlight_color']
        if original_formatting['strike'] is not None:
            run.font.strike = original_formatting['strike']
        if original_formatting['subscript'] is not None:
            run.font.subscript = original_formatting['subscript']
        if original_formatting['superscript'] is not None:
            run.font.superscript = original_formatting['superscript']
        if original_formatting['all_caps'] is not None:
            run.font.all_caps = original_formatting['all_caps']
        if original_formatting['small_caps'] is not None:
            run.font.small_caps = original_formatting['small_caps']

        if original_formatting['style_name']:
            try:
                run.style = original_formatting['style_name']
            except KeyError:
                print(f"Warning: Character style '{original_formatting['style_name']}' not found in document. Skipping style copy for run.")


    # Placeholder for style preservation mechanisms
    def copy_run_formatting(self, source_run, target_run):
        """
        Copies formatting from a source run to a target run.
        This includes common character formatting attributes.

        Args:
            source_run (Run): The run to copy formatting from.
            target_run (Run): The run to apply formatting to.
        """
        target_run.bold = source_run.bold
        target_run.italic = source_run.italic
        target_run.underline = source_run.underline
        target_run.font.name = source_run.font.name
        if source_run.font.size is not None:
            target_run.font.size = source_run.font.size
        if source_run.font.color and source_run.font.color.rgb is not None:
            target_run.font.color.rgb = source_run.font.color.rgb
        if source_run.font.highlight_color is not None:
            target_run.font.highlight_color = source_run.font.highlight_color
        target_run.font.strike = source_run.font.strike
        target_run.font.subscript = source_run.font.subscript
        target_run.font.superscript = source_run.font.superscript
        target_run.font.all_caps = source_run.font.all_caps
        target_run.font.small_caps = source_run.font.small_caps
        # Style name
        if source_run.style and source_run.style.name:
            try:
                target_run.style = source_run.style.name
            except KeyError:
                print(f"Warning: Character style '{source_run.style.name}' not found in document. Skipping style copy for run.")


    def copy_paragraph_formatting(self, source_paragraph, target_paragraph):
        """
        Copies formatting from a source paragraph to a target paragraph.
        This includes paragraph style and alignment.

        Args:
            source_paragraph (Paragraph): The paragraph to copy formatting from.
            target_paragraph (Paragraph): The paragraph to apply formatting to.
        """
        # Reason: Placeholder for future implementation.
        if source_paragraph.style and source_paragraph.style.name:
            try:
                target_paragraph.style = source_paragraph.style.name
            except KeyError:
                print(f"Warning: Paragraph style '{source_paragraph.style.name}' not found in document. Skipping style copy for paragraph.")
        
        if source_paragraph.alignment is not None:
            target_paragraph.alignment = source_paragraph.alignment
        
        # Paragraph format attributes (e.g., spacing, indentation)
        # These are part of ParagraphFormat object
        source_pf = source_paragraph.paragraph_format
        target_pf = target_paragraph.paragraph_format

        if source_pf.left_indent is not None:
            target_pf.left_indent = source_pf.left_indent
        if source_pf.right_indent is not None:
            target_pf.right_indent = source_pf.right_indent
        if source_pf.first_line_indent is not None:
            target_pf.first_line_indent = source_pf.first_line_indent
        if source_pf.space_before is not None:
            target_pf.space_before = source_pf.space_before
        if source_pf.space_after is not None:
            target_pf.space_after = source_pf.space_after
        if source_pf.line_spacing is not None:
            target_pf.line_spacing = source_pf.line_spacing
        if source_pf.line_spacing_rule is not None:
            target_pf.line_spacing_rule = source_pf.line_spacing_rule
        if source_pf.keep_together is not None:
            target_pf.keep_together = source_pf.keep_together
        if source_pf.keep_with_next is not None:
            target_pf.keep_with_next = source_pf.keep_with_next
        if source_pf.page_break_before is not None:
            target_pf.page_break_before = source_pf.page_break_before
        if source_pf.widow_control is not None:
            target_pf.widow_control = source_pf.widow_control
        
        # Tab stops (more complex, might need specific handling if required)
        # target_pf.tab_stops.clear_all()
        # for tab_stop in source_pf.tab_stops:
        #     target_pf.add_tab_stop(tab_stop.position, tab_stop.alignment, tab_stop.leader)

    def apply_changes(self, changes: List[DocumentChange]):
        """
        Applies a list of DocumentChange objects to the document.

        Args:
            changes (List[DocumentChange]): A list of changes to apply.

        Note: Applying changes that affect document structure (insertions, deletions)
        can affect the indices of subsequent changes. For simplicity, this
        implementation applies changes in the order provided. A more robust
        approach might require re-indexing or applying changes in reverse order
        for deletions.
        """
        for change in changes:
            try:
                if change.change_type == ChangeType.TEXT_UPDATE:
                    if isinstance(change.location, LocationParagraph):
                        para_index = change.location.paragraph_index
                        if 0 <= para_index < len(self.document.paragraphs):
                            paragraph = self.document.paragraphs[para_index]
                            if change.location.run_index is not None:
                                self._update_run_text(paragraph, change.location.run_index, str(change.new_value))
                            else:
                                # Apply to the whole paragraph, attempting to preserve style
                                self.update_paragraph_text(para_index, str(change.new_value), preserve_style=True)
                        else:
                            print(f"Warning: TEXT_UPDATE failed. Paragraph index {para_index} out of bounds.")
                    else:
                        print(f"Warning: TEXT_UPDATE requires LocationParagraph, but got {type(change.location).__name__}.")

                elif change.change_type == ChangeType.PARAGRAPH_INSERT:
                    if isinstance(change.location, LocationParagraph) and change.new_value is not None:
                         # Assuming location.paragraph_index is the index *before* which to insert
                        insert_before_index = change.location.paragraph_index
                        if 0 <= insert_before_index <= len(self.document.paragraphs):
                            # If inserting at the end, insert_paragraph_before on a non-existent index won't work.
                            # Need to handle appending specifically.
                            if insert_before_index == len(self.document.paragraphs):
                                self.document.add_paragraph(str(change.new_value))
                            else:
                                anchor_paragraph = self.document.paragraphs[insert_before_index]
                                anchor_paragraph.insert_paragraph_before(str(change.new_value))
                        else:
                             print(f"Warning: PARAGRAPH_INSERT failed. Insertion index {insert_before_index} out of bounds.")
                    else:
                        print(f"Warning: PARAGRAPH_INSERT requires LocationParagraph and new_value.")

                elif change.change_type == ChangeType.PARAGRAPH_DELETE:
                    if isinstance(change.location, LocationParagraph):
                        para_index = change.location.paragraph_index
                        self.delete_paragraph(para_index) # delete_paragraph handles index checks
                    else:
                        print(f"Warning: PARAGRAPH_DELETE requires LocationParagraph.")

                elif change.change_type == ChangeType.SECTION_REPLACE:
                    if isinstance(change.location, LocationSection) and isinstance(change.new_value, list):
                        self.replace_text_after_heading(
                            heading_text=change.location.heading_text,
                            new_content_paragraphs=change.new_value,
                            heading_style_name=change.location.heading_style_name
                        ) # replace_text_after_heading handles not finding the heading
                    else:
                        print(f"Warning: SECTION_REPLACE requires LocationSection and new_value as a list of strings.")

                elif change.change_type == ChangeType.TABLE_CELL_UPDATE:
                    if isinstance(change.location, LocationTable) and change.new_value is not None:
                        self.update_table_cell_text(
                            table_index=change.location.table_index,
                            row=change.location.row_index,
                            col=change.location.column_index,
                            text=str(change.new_value)
                        ) # update_table_cell_text handles index checks
                    else:
                        print(f"Warning: TABLE_CELL_UPDATE requires LocationTable and new_value.")

                # Add handling for other ChangeTypes as they are implemented

                else:
                    print(f"Warning: Unsupported change type: {change.change_type}")

                # Optionally update change status to APPLIED here if successful
                # change.status = ChangeStatus.APPLIED # This would require the changes list to be mutable or returning a new list

            except Exception as e:
                print(f"Error applying change {change.change_id}: {e}")
                # Optionally update change status to FAILED here
                # change.status = ChangeStatus.FAILED




    def get_content_between_headings(self, start_heading: str, end_heading: str) -> str:
        """
        Retrieves the content between two specified headings.

        Args:
            start_heading (str): The text of the starting heading.
            end_heading (str): The text of the ending heading.

        Returns:
            str: The content between the headings, or an empty string if the headings are not found.
        """
        content = ""
        start_found = False
        for para in self.document.paragraphs:
            if para.text == start_heading:
                start_found = True
            elif start_found and para.text == end_heading:
                break
            elif start_found:
                content += para.text + "\n"
        return content

if __name__ == '__main__':
    # Example Usage (for testing purposes)
    editor = WordEditor() # Creates a new document

    # Add a paragraph with specific formatting
    p1 = editor.document.add_paragraph()
    run1 = p1.add_run("This is a bold test.")
    run1.bold = True
    run1.font.name = 'Calibri'
    run1.font.size = Pt(14)

    p2 = editor.document.add_paragraph("This is a normal paragraph.")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save the new document
    new_doc_path = "test_document_new.docx"
    editor.save_document(new_doc_path)
    print(f"New document saved to {new_doc_path}")

    # Load an existing document (if you have one)
    # editor_existing = WordEditor("path_to_your_document.docx")
    # print(editor_existing.get_paragraph_text(0))