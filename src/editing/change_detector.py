# src/editing/change_detector.py
from typing import List, Tuple, Optional
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph as DocxParagraph

from src.editing.change_model import DocumentChange, ChangeType, LocationParagraph
from src.editing.word_editor import WordEditor # May be needed to load/process docs

class ChangeDetector:
    """
    Detects changes between two document states and generates DocumentChange objects.
    """

    def __init__(self, original_doc_path: str, modified_doc_path: str, document_id: Optional[str] = None):
        """
        Initializes the ChangeDetector with paths to the original and modified documents.

        Args:
            original_doc_path (str): Filepath to the original .docx document.
            modified_doc_path (str): Filepath to the modified .docx document.
            document_id (Optional[str]): A unique identifier for the document being compared.
                                         If None, the original_doc_path will be used.
        """
        self.original_doc_path = original_doc_path
        self.modified_doc_path = modified_doc_path
        self.document_id = document_id if document_id else original_doc_path

        # Load documents using WordEditor or directly with python-docx
        # For now, assume they are loaded elsewhere or loaded upon method call
        # self.original_docx = Document(original_doc_path)
        # self.modified_docx = Document(modified_doc_path)
        # Using WordEditor might be better if we need its helper functions
        self.original_editor = WordEditor(original_doc_path)
        self.modified_editor = WordEditor(modified_doc_path)


    def detect_paragraph_text_changes(self) -> List[DocumentChange]:
        """
        Detects text changes between paragraphs of the original and modified documents.
        This is a basic implementation focusing on 1-to-1 paragraph comparison by index.
        It does not yet handle inserted/deleted paragraphs robustly in this comparison.
        More advanced diffing (e.g., Myers diff) would be needed for that.

        Returns:
            List[DocumentChange]: A list of DocumentChange objects representing text updates.
        """
        changes: List[DocumentChange] = []
        
        original_paras = self.original_editor.document.paragraphs
        modified_paras = self.modified_editor.document.paragraphs

        # Simple comparison based on paragraph index
        # This won't correctly identify changes if paragraphs are inserted/deleted,
        # as indices will shift. A more sophisticated alignment is needed for that.
        len_orig = len(original_paras)
        len_mod = len(modified_paras)
        common_len = min(len_orig, len_mod)

        for i in range(common_len):
            orig_para_text = original_paras[i].text
            mod_para_text = modified_paras[i].text

            if orig_para_text != mod_para_text:
                # TODO: Implement run-level diffing within the paragraph for more granularity
                # For now, treating the whole paragraph text as changed.
                change = DocumentChange(
                    document_id=self.document_id,
                    change_type=ChangeType.TEXT_UPDATE,
                    location=LocationParagraph(paragraph_index=i),
                    old_value=orig_para_text,
                    new_value=mod_para_text,
                    description=f"Text changed in paragraph {i}."
                )
                changes.append(change)
        
        # Rudimentary detection of added/deleted paragraphs at the end
        if len_mod > len_orig:
            for i in range(len_orig, len_mod):
                change = DocumentChange(
                    document_id=self.document_id,
                    change_type=ChangeType.PARAGRAPH_INSERT, # Or treat as TEXT_UPDATE with old_value=None
                    location=LocationParagraph(paragraph_index=i),
                    old_value=None,
                    new_value=modified_paras[i].text,
                    description=f"Paragraph {i} inserted."
                )
                changes.append(change)
        elif len_orig > len_mod:
            for i in range(len_mod, len_orig):
                change = DocumentChange(
                    document_id=self.document_id,
                    change_type=ChangeType.PARAGRAPH_DELETE, # Or treat as TEXT_UPDATE with new_value=None
                    location=LocationParagraph(paragraph_index=i),
                    old_value=original_paras[i].text,
                    new_value=None,
                    description=f"Paragraph {i} deleted."
                )
                changes.append(change)
                
        return changes

    # Placeholder for detecting formatting changes
    def detect_formatting_changes(self) -> List[DocumentChange]:
        """Placeholder for detecting formatting changes."""
        # Reason: Placeholder for future implementation.
        return []

    # Placeholder for detecting structural changes (insertions, deletions beyond simple tail)
    def detect_structural_changes(self) -> List[DocumentChange]:
        """Placeholder for detecting structural changes like insertions/deletions."""
        # Reason: Placeholder for future implementation.
        return []

    def detect_all_changes(self) -> List[DocumentChange]:
        """Detects all types of changes."""
        all_changes: List[DocumentChange] = []
        all_changes.extend(self.detect_paragraph_text_changes())
        all_changes.extend(self.detect_formatting_changes())
        all_changes.extend(self.detect_structural_changes())
        # TODO: Add table change detection, list change detection etc.
        return all_changes

if __name__ == '__main__':
    # This is example usage and requires actual docx files.
    # Create dummy docx files for a basic test.
    
    # Create original_doc.docx
    doc_orig = DocxDocument()
    doc_orig.add_paragraph("This is the first paragraph, unchanged.")
    doc_orig.add_paragraph("This is the second paragraph, which will be modified.")
    doc_orig.add_paragraph("This paragraph will be deleted in the modified version.")
    doc_orig.save("original_doc.docx")

    # Create modified_doc.docx
    doc_mod = DocxDocument()
    doc_mod.add_paragraph("This is the first paragraph, unchanged.")
    doc_mod.add_paragraph("This is the second paragraph, now it is modified.") # Changed text
    # Third paragraph is omitted (deleted)
    doc_mod.add_paragraph("This is a new fourth paragraph, inserted at the end.") # Inserted
    doc_mod.save("modified_doc.docx")

    print("Created dummy original_doc.docx and modified_doc.docx for testing.")
    print("Please run tests or use this detector with these files.")

    # Example of using the detector:
    # detector = ChangeDetector("original_doc.docx", "modified_doc.docx", document_id="test_doc_comparison")
    # text_changes = detector.detect_paragraph_text_changes()
    # print("\nDetected Text Changes:")
    # for change in text_changes:
    #     print(change.model_dump_json(indent=2))
    
    # To clean up:
    # import os
    # os.remove("original_doc.docx")
    # os.remove("modified_doc.docx")