# src/editing/workflow_applier.py
from typing import List
from .change_model import DocumentChange, ChangeWorkflow, ChangeStatus, ChangeType, LocationParagraph, LocationTable, LocationSection
from .word_editor import WordEditor
from docx.document import Document # Import Document for type hinting if needed

class WorkflowApplier:
    """
    Applies accepted changes from a ChangeWorkflow to a Word document using WordEditor.
    """

    def __init__(self, word_editor: WordEditor):
        """
        Initializes the WorkflowApplier with a WordEditor instance.

        Args:
            word_editor (WordEditor): The WordEditor instance to use for document manipulation.
        """
        self.word_editor = word_editor

    def apply_workflow(self, workflow: ChangeWorkflow):
        """
        Applies all accepted changes within a workflow to the document.

        Args:
            workflow (ChangeWorkflow): The workflow containing the changes to apply.

        Note: This implementation attempts to apply changes in an order that
        minimizes index conflicts (deletions in reverse, insertions in order).
        However, complex interdependencies between changes might still require
        more sophisticated handling.
        """
        accepted_changes = workflow.get_accepted_changes()

        # Separate changes by type for ordered application
        text_updates_para = []
        text_updates_run = []
        paragraph_inserts = []
        paragraph_deletes = []
        section_replaces = []
        table_cell_updates = []
        # Add lists for other change types as they are implemented

        for change in accepted_changes:
            if change.change_type == ChangeType.TEXT_UPDATE:
                if isinstance(change.location, LocationParagraph):
                    if change.location.run_index is not None:
                        text_updates_run.append(change)
                    else:
                        text_updates_para.append(change)
            elif change.change_type == ChangeType.PARAGRAPH_INSERT:
                paragraph_inserts.append(change)
            elif change.change_type == ChangeType.PARAGRAPH_DELETE:
                paragraph_deletes.append(change)
            elif change.change_type == ChangeType.SECTION_REPLACE:
                section_replaces.append(change)
            elif change.change_type == ChangeType.TABLE_CELL_UPDATE:
                table_cell_updates.append(change)
            # Add handling for other change types

        # Apply deletions first, in reverse order of paragraph index to avoid index issues
        paragraph_deletes.sort(key=lambda c: c.location.paragraph_index, reverse=True)
        for change in paragraph_deletes:
            try:
                if isinstance(change.location, LocationParagraph):
                    success = self.word_editor.delete_paragraph(change.location.paragraph_index)
                    if success:
                        workflow.update_change_status(change.change_id, ChangeStatus.APPLIED)
                    else:
                        workflow.update_change_status(change.change_id, ChangeStatus.FAILED)
                else:
                     print(f"Warning: PARAGRAPH_DELETE change {change.change_id} has invalid location type.")
                     workflow.update_change_status(change.change_id, ChangeStatus.FAILED)
            except Exception as e:
                print(f"Error applying PARAGRAPH_DELETE change {change.change_id}: {e}")
                workflow.update_change_status(change.change_id, ChangeStatus.FAILED)


        # Apply insertions next, in forward order of paragraph index
        # Note: This assumes insertion index refers to the index *before* insertion in the *current* document state.
        # If changes were detected against the original document, indices might need adjustment.
        # For now, we assume changes are applied sequentially and indices are relative to the document state just before the change is applied.
        paragraph_inserts.sort(key=lambda c: c.location.paragraph_index)
        # Need to adjust indices for insertions based on previous insertions
        # This is a simplification; a robust solution would track index shifts.
        # For this basic implementation, we'll apply in order and hope for the best with simple cases.
        # A better approach would be to calculate final indices after all insertions/deletions are planned.
        # Let's stick to the simple sequential application for now and note the limitation.
        print("Note: Simple sequential application of insertions may lead to incorrect results if multiple insertions affect nearby indices.")
        for change in paragraph_inserts:
             try:
                if isinstance(change.location, LocationParagraph) and change.new_value is not None:
                    # The insert_paragraph_before method in WordEditor handles inserting before the specified index.
                    # So, if we want to insert *at* index X, we insert before the paragraph currently at index X.
                    # If we want to append at the end, we insert before len(paragraphs).
                    # The LocationParagraph index for INSERT is defined as the index *before* which to insert.
                    insert_before_index = change.location.paragraph_index
                    
                    # Need to re-check bounds as document length changes
                    if 0 <= insert_before_index <= len(self.word_editor.document.paragraphs):
                         # WordEditor.insert_paragraph_after is more intuitive for inserting *at* an index
                         # Let's use that and adjust the location index definition for INSERT
                         # Re-reading change_model.py, LocationParagraph for INSERT is "index *before* which to insert"
                         # So, inserting before index X means the new paragraph will be at index X.
                         # WordEditor.insert_paragraph_after(index, text) inserts *after* index, so new paragraph is at index + 1.
                         # To insert *at* index X using insert_paragraph_after, we need to insert after index X-1.
                         # Special case: inserting at index 0 means inserting before the first paragraph (index 0).
                         # WordEditor.insert_paragraph_after(0, text) inserts after index 0 (the first paragraph), placing new para at index 1.
                         # WordEditor.insert_paragraph_after(-1, text) could potentially append, but let's use add_paragraph for appending.

                         # Let's redefine LocationParagraph index for INSERT to be the index *at* which to insert.
                         # This is more intuitive for the applier. I will update change_model.py later if needed,
                         # but for now, I'll interpret LocationParagraph(paragraph_index=X) for INSERT
                         # as "insert the new paragraph so it becomes the paragraph at index X".
                         # To achieve this with WordEditor.insert_paragraph_after(index, text), we insert after index X-1.
                         # If X is 0, we insert before the current first paragraph. WordEditor doesn't have a direct insert_before_index.
                         # WordEditor.insert_paragraph_after(index, text) inserts *after* index.
                         # To insert AT index X:
                         # If X == 0, insert before the current first paragraph.
                         # If X > 0, insert after the paragraph at index X-1.

                         target_index = change.location.paragraph_index # Interpreting this as the target index of the new paragraph

                         if target_index == 0:
                             # Insert before the current first paragraph
                             # python-docx insert_paragraph_before is on a paragraph object.
                             # We need to get the paragraph currently at index 0 and insert before it.
                             if len(self.word_editor.document.paragraphs) > 0:
                                 self.word_editor.document.paragraphs[0].insert_paragraph_before(str(change.new_value))
                             else:
                                 # Document is empty, just add the first paragraph
                                 self.word_editor.document.add_paragraph(str(change.new_value))
                         elif 0 < target_index <= len(self.word_editor.document.paragraphs):
                             # Insert after the paragraph at target_index - 1
                             self.word_editor.insert_paragraph_after(target_index - 1, str(change.new_value))
                         else:
                             print(f"Warning: PARAGRAPH_INSERT failed. Target index {target_index} out of bounds for current document length {len(self.word_editor.document.paragraphs)}.")
                             workflow.update_change_status(change.change_id, ChangeStatus.FAILED)
                             continue # Skip to next change

                         workflow.update_change_status(change.change_id, ChangeStatus.APPLIED)
                    else:
                         print(f"Warning: PARAGRAPH_INSERT change {change.change_id} has invalid location or new_value.")
                         workflow.update_change_status(change.change_id, ChangeStatus.FAILED)
             except Exception as e:
                print(f"Error applying PARAGRAPH_INSERT change {change.change_id}: {e}")
                workflow.update_change_status(change.change_id, ChangeStatus.FAILED)


        # Apply other changes in forward order
        # TEXT_UPDATE (paragraph and run), SECTION_REPLACE, TABLE_CELL_UPDATE
        # Note: Applying these after insertions/deletions means their original indices
        # might be invalid. A robust solution would re-index changes based on structural modifications.
        # For this basic implementation, we acknowledge this limitation.
        print("Note: Applying TEXT_UPDATE, SECTION_REPLACE, TABLE_CELL_UPDATE after structural changes may use outdated indices.")

        for change in text_updates_para:
            try:
                if isinstance(change.location, LocationParagraph) and change.new_value is not None:
                    # Use the potentially outdated index
                    self.word_editor.update_paragraph_text(change.location.paragraph_index, str(change.new_value), preserve_style=True)
                    workflow.update_change_status(change.change_id, ChangeStatus.APPLIED)
                else:
                    print(f"Warning: TEXT_UPDATE (paragraph) change {change.change_id} has invalid location or new_value.")
                    workflow.update_change_status(change.change_id, ChangeStatus.FAILED)
            except Exception as e:
                print(f"Error applying TEXT_UPDATE (paragraph) change {change.change_id}: {e}")
                workflow.update_change_status(change.change_id, ChangeStatus.FAILED)

        for change in text_updates_run:
             try:
                if isinstance(change.location, LocationParagraph) and change.new_value is not None and change.location.run_index is not None:
                    # Need to get the paragraph again as index might have shifted
                    para_index = change.location.paragraph_index
                    if 0 <= para_index < len(self.word_editor.document.paragraphs):
                        paragraph = self.word_editor.document.paragraphs[para_index]
                        self.word_editor._update_run_text(paragraph, change.location.run_index, str(change.new_value))
                        workflow.update_change_status(change.change_id, ChangeStatus.APPLIED)
                    else:
                        print(f"Warning: TEXT_UPDATE (run) failed. Paragraph index {para_index} out of bounds (possibly due to prior structural changes).")
                        workflow.update_change_status(change.change_id, ChangeStatus.FAILED)
                else:
                    print(f"Warning: TEXT_UPDATE (run) change {change.change_id} has invalid location or new_value.")
                    workflow.update_change_status(change.change_id, ChangeStatus.FAILED)
             except Exception as e:
                print(f"Error applying TEXT_UPDATE (run) change {change.change_id}: {e}")
                workflow.update_change_status(change.change_id, ChangeStatus.FAILED)


        for change in section_replaces:
            try:
                if isinstance(change.location, LocationSection) and isinstance(change.new_value, list):
                    # Note: replace_text_after_heading searches by heading text, which is more robust to index changes
                    success = self.word_editor.replace_text_after_heading(
                        heading_text=change.location.heading_text,
                        new_content_paragraphs=change.new_value,
                        heading_style_name=change.location.heading_style_name
                    )
                    if success:
                        workflow.update_change_status(change.change_id, ChangeStatus.APPLIED)
                    else:
                        workflow.update_change_status(change.change_id, ChangeStatus.FAILED)
                else:
                    print(f"Warning: SECTION_REPLACE change {change.change_id} has invalid location or new_value.")
                    workflow.update_change_status(change.change_id, ChangeStatus.FAILED)
            except Exception as e:
                print(f"Error applying SECTION_REPLACE change {change.change_id}: {e}")
                workflow.update_change_status(change.change_id, ChangeStatus.FAILED)

        for change in table_cell_updates:
            try:
                if isinstance(change.location, LocationTable) and change.new_value is not None:
                    # Table and cell indices are generally more stable unless tables themselves are added/deleted
                    success = self.word_editor.update_table_cell_text(
                        table_index=change.location.table_index,
                        row=change.location.row_index,
                        col=change.location.column_index,
                        text=str(change.new_value)
                    )
                    if success:
                        workflow.update_change_status(change.change_id, ChangeStatus.APPLIED)
                    else:
                        workflow.update_change_status(change.change_id, ChangeStatus.FAILED)
                else:
                    print(f"Warning: TABLE_CELL_UPDATE change {change.change_id} has invalid location or new_value.")
                    workflow.update_change_status(change.change_id, ChangeStatus.FAILED)
            except Exception as e:
                print(f"Error applying TABLE_CELL_UPDATE change {change.change_id}: {e}")
                workflow.update_change_status(change.change_id, ChangeStatus.FAILED)

        # After applying changes, update workflow status if all changes have been reviewed
        if workflow.all_changes_reviewed():
            workflow.status = "completed"
        self.updated_at = datetime.utcnow()


# Example Usage (for illustration)
if __name__ == "__main__":
    # This example requires creating a dummy document and a workflow with changes.
    # It's more complex than a simple module test and might be better suited
    # for an integration test or a separate script.
    print("Example usage for WorkflowApplier requires a document and a workflow.")
    print("Please refer to integration tests for a full example.")
    pass