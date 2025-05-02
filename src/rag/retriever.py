import os
import json
from typing import List, Dict, Any, Optional
from rank_bm25 import BM25Okapi
from collections import defaultdict

from .embeddings import EmbeddingModel
from .vector_store import VectorStore
from ..data.metadata_store import MetadataStore
from ..llm.gemini import GeminiClient  # Import GeminiClient
from ..llm.prompts import PromptTemplates  # Import PromptTemplates
from ..config import settings  # Assuming config is in src/


# Simple tokenizer for BM25
def simple_tokenizer(text: str) -> List[str]:
    return text.lower().split()


def estimate_token_count(text: str) -> int:
    """Estimates the token count of a given text using a simple word-based approximation."""
    return len(simple_tokenizer(text))


class Retriever:
    """
    Handles retrieving relevant document chunks based on a query and generating responses.
    Supports vector search and hybrid search (BM25 + Vector) using RRF.
    Uses the EmbeddingModel for query embedding, VectorStore for similarity search,
    MetadataStore to retrieve full chunk data, and GeminiClient for response generation.
    """

    def __init__(
        self,
        embedding_model: EmbeddingModel,
        vector_store: VectorStore,
        metadata_store: MetadataStore,
        gemini_client: GeminiClient,
    ):  # Add GeminiClient
        """
        Initializes the Retriever.

        Args:
            embedding_model: An instance of the EmbeddingModel.
            vector_store: An instance of the VectorStore.
            metadata_store: An instance of the MetadataStore.
            gemini_client: An instance of the GeminiClient.
        """
        self.embedding_model = embedding_model
        self.vector_store = vector_store
        self.metadata_store = metadata_store
        self.gemini_client = gemini_client
        self.bm25: Optional[BM25Okapi] = None
        self.chunk_id_corpus_map: Optional[Dict[int, str]] = None  # Map chunk_id to content for BM25 retrieval
        self.retrieval_cache: Dict[str, List[Dict[str, Any]]] = {}  # Cache retrieval results

    def build_bm25_index(self):
        """
        Builds the BM25 index from all chunks stored in the MetadataStore.
        This should be called after indexing documents and before retrieval.
        """
        print("Building BM25 index...")
        # Fetch all chunks (id and content) from the metadata store
        # This might be memory intensive for very large datasets!
        all_chunks = []
        # Assuming we need to iterate through all documents first
        # This part needs refinement based on how MetadataStore can efficiently retrieve all chunks
        # For now, let's assume a hypothetical `get_all_chunks` method exists
        # Replace this with actual implementation if MetadataStore is updated
        try:
            # A more direct way if possible:
            conn = self.metadata_store._get_connection()
            cursor = conn.cursor()
            cursor.execute("SELECT id, content FROM chunks")  # Select only id and content
            all_chunks_raw = cursor.fetchall()
            conn.close()
            all_chunks = [{"id": row[0], "content": row[1]} for row in all_chunks_raw]

        except Exception as e:
            print(f"Error fetching all chunks for BM25 index: {e}")
            return

        if not all_chunks:
            print("No chunks found in metadata store to build BM25 index.")
            return

        self.chunk_id_corpus_map = {chunk['id']: chunk['content'] for chunk in all_chunks}
        corpus_texts = [chunk['content'] for chunk in all_chunks]
        tokenized_corpus = [simple_tokenizer(doc) for doc in corpus_texts]

        self.bm25 = BM25Okapi(tokenized_corpus)
        print(f"BM25 index built with {len(corpus_texts)} chunks.")

    def retrieve(
        self,
        query: str,
        k: int = 5,
        search_type: str = "hybrid",
        rrf_k: int = 60,
        metadata_filters: Optional[Dict[str, Any]] = None,
    ) -> List[Dict[str, Any]]:
        """
        Retrieves the top-k most relevant document chunks for a given query and generates a response.
        Caches retrieval results to improve performance.

        Args:
            query: The search query string.
            k: The number of relevant chunks to retrieve.
            search_type: Type of search ('vector', 'bm25', 'hybrid'). Default is 'hybrid'.
            rrf_k: The constant used in Reciprocal Rank Fusion (RRF) calculation. Default is 60.
            metadata_filters: Optional dictionary containing metadata fields and values to filter by.

        Returns:
            A list of dictionaries, where each dictionary contains the data
            (including content and metadata) of a retrieved chunk, along with
            its final score or distance, and the generated response.
        """
        if not query:
            return []

        # Create a cache key
        cache_key = json.dumps(
            {
                "query": query,
                "k": k,
                "search_type": search_type,
                "rrf_k": rrf_k,
                "metadata_filters": metadata_filters,
            },
            sort_keys=True,
        )

        # Check if the result is already cached
        if cache_key in self.retrieval_cache:
            print("Retrieving results from cache - HIT")
            return self.retrieval_cache[cache_key]
        else:
            print("Retrieving results from cache - MISS")

        vector_results = []
        bm25_results = []
        final_results_ids = set()

        # 1. Vector Search (if applicable)
        if search_type in ["vector", "hybrid"]:
            print(f"Performing vector search for query: '{query}'")
            query_embedding = self.embedding_model.generate_embeddings([query])[0]
            # Retrieve more results initially for potential re-ranking
            initial_k_vector = k * 2 if search_type == "hybrid" else k
            vector_search_raw = self.vector_store.search(query_embedding, k=initial_k_vector)

            # Apply metadata filters
            filtered_vector_search_raw = []
            for res in vector_search_raw:
                chunk_id = int(res['external_id'])
                chunk_data = self.metadata_store.get_chunk_by_id(chunk_id)
                if chunk_data:
                    metadata = json.loads(chunk_data.get("metadata", "{}"))
                    if self._matches_filters(metadata, metadata_filters):
                        filtered_vector_search_raw.append(res)

            # Assuming external_id is chunk_id (needs update in indexer)
            vector_results = [(int(res['external_id']), 1.0 / (i + 1), res['distance']) for i, res in enumerate(filtered_vector_search_raw)]  # (chunk_id, rank_score, distance)
            final_results_ids.update([res[0] for res in vector_results])
            print(f"Vector search found {len(vector_results)} results.")

        # 2. BM25 Search (if applicable)
        if search_type in ["bm25", "hybrid"]:
            if self.bm25 is None or self.chunk_id_corpus_map is None:
                print("Warning: BM25 index not built. Building now...")
                self.build_bm25_index()  # Attempt to build if not already built

            if self.bm25 and self.chunk_id_corpus_map:
                print(f"Performing BM25 search for query: '{query}'")
                tokenized_query = simple_tokenizer(query)
                # Get scores for all documents, then sort and take top N
                # Note: BM25Okapi.get_scores returns scores in the original corpus order
                # We need to map these scores back to chunk_ids and sort
                doc_scores = self.bm25.get_scores(tokenized_query)

                # Apply metadata filters
                filtered_chunk_id_scores = {}
                corpus_index_to_chunk_id = {i: chunk_id for i, chunk_id in enumerate(self.chunk_id_corpus_map.keys())}
                for i, score in enumerate(doc_scores):
                    chunk_id = corpus_index_to_chunk_id[i]
                    chunk_data = self.metadata_store.get_chunk_by_id(chunk_id)
                    if chunk_data:
                        metadata = json.loads(chunk_data.get("metadata", "{}"))
                        if self._matches_filters(metadata, metadata_filters):
                            filtered_chunk_id_scores[chunk_id] = score

                # Sort by score descending
                sorted_bm25 = sorted(filtered_chunk_id_scores.items(), key=lambda item: item[1], reverse=True)

                # Retrieve more results initially for potential re-ranking
                initial_k_bm25 = k * 2 if search_type == "hybrid" else k
                bm25_results = [(chunk_id, 1.0 / (i + 1), score) for i, (chunk_id, score) in enumerate(sorted_bm25[:initial_k_bm25])]  # (chunk_id, rank_score, bm25_score)
                final_results_ids.update([res[0] for res in bm25_results])
                print(f"BM25 search found {len(bm25_results)} results.")
            else:
                print("BM25 index not available, skipping BM25 search.")

        # 3. Combine and Re-rank (if hybrid)
        combined_results = []
        if search_type == "hybrid":
            print("Combining results using Reciprocal Rank Fusion (RRF)...")
            rrf_scores = defaultdict(float)
            # Process vector results
            for rank, (chunk_id, rank_score, distance) in enumerate(vector_results):
                rrf_scores[chunk_id] += 1.0 / (rrf_k + rank + 1)  # RRF formula

            # Process BM25 results
            for rank, (chunk_id, rank_score, bm25_score) in enumerate(bm25_results):
                rrf_scores[chunk_id] += 1.0 / (rrf_k + rank + 1)  # RRF formula

            # ----------------------------------------------------------------------------------------------------
            # Add chunk length to the RRF score
            # ----------------------------------------------------------------------------------------------------
            for chunk_id in rrf_scores:
                chunk_data = self.metadata_store.get_chunk_by_id(chunk_id)
                if chunk_data:
                    chunk_length = len(chunk_data["content"])
                    # Normalize chunk length (assuming a reasonable range, e.g., 50-200 tokens)
                    normalized_length = max(0, min(1, (chunk_length - 50) / 150))
                    rrf_scores[chunk_id] += 0.1 * normalized_length  # Adjust the weight (0.1) as needed

            # ----------------------------------------------------------------------------------------------------
            # Add metadata similarity to the RRF score
            # ----------------------------------------------------------------------------------------------------
            # For simplicity, let's assume we have a 'topic' metadata field and boost scores
            # for chunks that have the same topic as the query (if specified)
            # This requires a way to extract topic from the query (e.g., using another LLM call)
            # For now, let's assume the query itself is the topic
            query_topic = query.lower()  # Treat the query as the topic
            for chunk_id in rrf_scores:
                chunk_data = self.metadata_store.get_chunk_by_id(chunk_id)
                if chunk_data:
                    metadata = json.loads(chunk_data.get("metadata", "{}"))
                    chunk_topic = metadata.get("topic", "").lower()
                    if chunk_topic == query_topic:
                        rrf_scores[chunk_id] += 0.2  # Adjust the weight (0.2) as needed

            # ----------------------------------------------------------------------------------------------------
            # Add positional context to the RRF score
            # ----------------------------------------------------------------------------------------------------
            # Boost scores for chunks that appear earlier in the document
            for chunk_id in rrf_scores:
                chunk_data = self.metadata_store.get_chunk_by_id(chunk_id)
                if chunk_data:
                    chunk_index = chunk_data.get("chunk_index", 0)
                    # Normalize chunk index (assuming chunks are roughly in order)
                    normalized_index = max(0, min(1, 1.0 / (chunk_index + 1)))  # Invert the index
                    rrf_scores[chunk_id] += 0.05 * normalized_index  # Adjust the weight (0.05) as needed

            # Sort combined results by RRF score
            sorted_combined_ids = sorted(rrf_scores.keys(), key=lambda cid: rrf_scores[cid], reverse=True)
            final_results_ids = sorted_combined_ids[:k]  # Get top k combined results

        elif search_type == "vector":
            final_results_ids = [res[0] for res in sorted(vector_results, key=lambda x: x[2])[:k]]  # Sort by distance
        elif search_type == "bm25":
            final_results_ids = [res[0] for res in sorted(bm25_results, key=lambda x: x[2], reverse=True)[:k]]  # Sort by score

        # 4. Retrieve full chunk data for final results
        retrieved_chunks_data: List[Dict[str, Any]] = []
        print(f"Retrieving final {len(final_results_ids)} chunks from metadata store...")
        for chunk_id in final_results_ids:
            chunk_data = self.metadata_store.get_chunk_by_id(chunk_id)  # Use get_chunk_by_id
            if chunk_data:
                # Optionally add the final score (RRF, distance, or BM25)
                if search_type == "hybrid":
                    chunk_data["score"] = rrf_scores.get(chunk_id)
                elif search_type == "vector":
                    # Find the original vector result for this chunk_id
                    vec_res = next((vr for vr in vector_results if vr[0] == chunk_id), None)
                    chunk_data["distance"] = vec_res[2] if vec_res else None
                elif search_type == "bm25":
                    # Find the original bm25 result for this chunk_id
                    bm25_res = next((br for br in bm25_results if br[0] == chunk_id), None)
                    chunk_data["score"] = bm25_res[2] if bm25_res else None

                retrieved_chunks_data.append(chunk_data)
            else:
                print(f"Warning: Chunk data not found for final chunk_id {chunk_id}")

        # Sort final results based on search type
        if search_type == "hybrid":
            retrieved_chunks_data.sort(key=lambda x: x.get("score", 0), reverse=True)
        elif search_type == "vector":
            retrieved_chunks_data.sort(key=lambda x: x.get("distance", float('inf')))
        elif search_type == "bm25":
            retrieved_chunks_data.sort(key=lambda x: x.get("score", 0), reverse=True)

        # 5. Context Window Management
        print("Managing context window...")
        max_context_tokens = 4096  # Example token limit for Gemini
        context = ""
        selected_chunks = []
        current_token_count = 0

        for chunk in retrieved_chunks_data:
            chunk_content = chunk["content"]
            chunk_tokens = estimate_token_count(chunk_content)

            if current_token_count + chunk_tokens <= max_context_tokens:
                context += chunk_content + "\n"
                selected_chunks.append(chunk)
                current_token_count += chunk_tokens
            else:
                # Truncate the chunk if it exceeds the remaining token limit
                remaining_tokens = max_context_tokens - current_token_count
                if remaining_tokens > 0:
                    truncated_chunk = chunk_content[: remaining_tokens * 4]  # Assuming 1 token is roughly 4 characters
                    context += truncated_chunk + "\n"
                    selected_chunks.append(chunk)
                    current_token_count = max_context_tokens  # Set to max
                break  # Stop adding chunks

        print(f"Selected {len(selected_chunks)} chunks for context. Estimated token count: {current_token_count}")

        # 6. Generate response using Gemini
        print("Generating response using Gemini...")
        # Combine the content of the selected chunks into a single context string
        # Create the prompt using the QA template
        prompt = PromptTemplates.qa_template(context=context, query=query)
        # Generate the response
        response_text = self.gemini_client.generate_response(prompt)

        # Add the response to the result
        results = []
        for chunk in selected_chunks:
            results.append({**chunk, "response": response_text})  # Merge chunk data with response

        print(f"Retrieved {len(results)} final chunks and generated response.")
        
        # Store the result in the cache
        self.retrieval_cache[cache_key] = results
        return results

    def _matches_filters(self, metadata: Dict[str, Any], metadata_filters: Optional[Dict[str, Any]]) -> bool:
        """
        Checks if a chunk's metadata matches the specified filters.
        """
        if not metadata_filters:
            return True

        for field, value in metadata_filters.items():
            if field not in metadata or metadata[field] != value:
                return False

        return True

    def clear_cache(self):
        """Clears the retrieval cache."""
        self.retrieval_cache = {}
        print("Retrieval cache cleared.")

    def invalidate_cache(self):
        """Invalidates the retrieval cache when BM25 index is rebuilt or new documents are added."""
        self.clear_cache()

# Example Usage (Needs significant updates to reflect changes)
if __name__ == "__main__":
    # This example requires initialized components and indexed data
    # Ensure you have run the indexing example first to create dummy data and index
    # CRITICAL: The indexing example needs to be updated to store chunk_id as external_id in VectorStore
    print("NOTE: The example usage below requires the DocumentIndexer to be updated")
    print("      to store metadata chunk IDs as external_ids in the VectorStore.")
    print("      Running this example without that update might lead to errors.")

    test_vector_db_dir = "data/vector_db_test_retrieval"  # Use a separate dir for this test
    test_metadata_db_path = "data/metadata_test_retrieval/metadata.db"  # Use a separate db

    # Ensure dummy files and directories exist
    dummy_docx_path = "data/documents/dummy_retrieval.docx"
    os.makedirs("data/documents", exist_ok=True)
    os.makedirs(os.path.dirname(test_metadata_db_path), exist_ok=True)
    os.makedirs(test_vector_db_dir, exist_ok=True)

    # Create dummy file if it doesn't exist
    if not os.path.exists(dummy_docx_path):
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument()
            doc.add_paragraph("Hybrid search combines vector similarity with keyword matching.")
            doc.add_paragraph("BM25 is a popular keyword-based algorithm.")
            doc.add_paragraph("Vector search finds semantically similar text.")
            doc.save(dummy_docx_path)
            print(f"Created dummy docx file: {dummy_docx_path}")
        except ImportError:
            print("python-docx not installed. Cannot create dummy docx.")
        except Exception as e:
            print(f"Error creating dummy docx: {e}")

    try:
        # Initialize components
        print("\nInitializing components for retrieval test...")
        embedding_model = EmbeddingModel()
        vector_store = VectorStore(index_path=test_vector_db_dir)
        metadata_store = MetadataStore(db_path=test_metadata_db_path)
        gemini_client = GeminiClient()  # Initialize GeminiClient

        # Index the document (assuming indexer is updated to store chunk_id)
        print("\nRunning indexing...")
        from .indexing import DocumentIndexer  # Re-import for clarity
        # THIS WILL FAIL if indexer is not updated to store chunk_id
        document_indexer = DocumentIndexer(vector_store, embedding_model, metadata_store, chunk_size=50, chunk_overlap=10)
        if os.path.exists(dummy_docx_path):
            document_indexer.index_document(dummy_docx_path)
        vector_store.save()  # Save after indexing

        # Initialize Retriever and build BM25 index
        retriever = Retriever(embedding_model, vector_store, metadata_store, gemini_client)
        retriever.build_bm25_index()  # Build the index

        # Perform retrievals
        query = "keyword algorithms"
        print(f"\nPerforming HYBRID retrieval for query: '{query}'")
        hybrid_chunks = retriever.retrieve(query, k=2, search_type="hybrid")

        print("\nHybrid Retrieved Chunks:")
        if hybrid_chunks:
            for i, chunk in enumerate(hybrid_chunks):
                print(f"--- Chunk {i+1} (Score: {chunk.get('score', 0):.4f}) ---")
                print(f"Content: {chunk.get('content', '')}")
                print(f"Response: {chunk.get('response', 'No response generated')}")  # Print the response
        else:
            print("No chunks retrieved.")

        print(f"\nPerforming VECTOR retrieval for query: '{query}'")
        vector_chunks = retriever.retrieve(query, k=2, search_type="vector")
        print("\nVector Retrieved Chunks:")
        if vector_chunks:
            for i, chunk in enumerate(vector_chunks):
                print(f"--- Chunk {i+1} (Distance: {chunk.get('distance', float('inf')):.4f}) ---")
                print(f"Content: {chunk.get('content', '')}")
                print(f"Response: {chunk.get('response', 'No response generated')}")  # Print the response
        else:
            print("No chunks retrieved.")

        print(f"\nPerforming BM25 retrieval for query: '{query}'")
        bm25_chunks = retriever.retrieve(query, k=2, search_type="bm25")
        print("\nBM25 Retrieved Chunks:")
        if bm25_chunks:
            for i, chunk in enumerate(bm25_chunks):
                print(f"--- Chunk {i+1} (Score: {chunk.get('score', 0):.4f}) ---")
                print(f"Content: {chunk.get('content', '')}")
                print(f"Response: {chunk.get('response', 'No response generated')}")  # Print the response
        else:
            print("No chunks retrieved.")

        # Example of filtering by metadata
        print(f"\nPerforming HYBRID retrieval with metadata filtering (topic='SOP')...")
        filtered_chunks = retriever.retrieve(query, k=2, search_type="hybrid", metadata_filters={"topic": "SOP"})

        print("\nFiltered Retrieved Chunks:")
        if filtered_chunks:
            for i, chunk in enumerate(filtered_chunks):
                print(f"--- Chunk {i+1} (Score: {chunk.get('score', 0):.4f}) ---")
                print(f"Content: {chunk.get('content', '')}")
                print(f"Metadata: {json.dumps(chunk.get('metadata', {}))}")
                print(f"Response: {chunk.get('response', 'No response generated')}")  # Print the response
        else:
            print("No chunks retrieved.")

        # Example of clearing the cache
        print("\nClearing the retrieval cache...")
        retriever.clear_cache()
        print("Retrieval cache cleared.")

    except ImportError as e:
        print(f"Could not import necessary modules for retrieval test: {e}. Skipping test.")
    except Exception as e:
        print(f"Error during retrieval test: {e}")
    finally:
        # Clean up
        if os.path.exists(dummy_docx_path):
            os.remove(dummy_docx_path)
        if os.path.exists(test_vector_db_dir):
            import shutil
            shutil.rmtree(test_vector_db_dir)
        if os.path.exists(test_metadata_db_path):
            os.remove(test_metadata_db_path)
        if os.path.exists(os.path.dirname(test_metadata_db_path)):
            try:
                os.rmdir(os.path.dirname(test_metadata_db_path))
            except OSError:
                pass
        print(f"\nCleaned up test files and directories.")