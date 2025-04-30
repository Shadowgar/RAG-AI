import os
import time
from datetime import datetime # Import datetime for file metadata
from typing import List, Dict, Any
from ..processors.document_parser import parse_document
from ..processors.chunking import chunk_document_elements
from .embeddings import EmbeddingModel
from .vector_store import VectorStore
from ..data.metadata_store import MetadataStore # Import MetadataStore
from ..config import settings

class DocumentIndexer:
    """
    Orchestrates the document indexing pipeline: parse -> chunk -> metadata -> embed -> store.
    """
    def __init__(self,
                 vector_store: VectorStore,
                 embedding_model: EmbeddingModel,
                 metadata_store: MetadataStore, # Add MetadataStore parameter
                 chunk_size: int = settings.EMBEDDING_BATCH_SIZE,
                 chunk_overlap: int = 50):
        """
        Initializes the DocumentIndexer.

        Args:
            vector_store: An instance of the VectorStore.
            embedding_model: An instance of the EmbeddingModel.
            metadata_store: An instance of the MetadataStore.
            chunk_size: The size of the document chunks.
            chunk_overlap: The overlap between document chunks.
        """
        self.vector_store = vector_store
        self.embedding_model = embedding_model
        self.metadata_store = metadata_store # Store the metadata store instance
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def index_document(self, file_path: str):
        """
        Processes and indexes a single document file, including storing metadata
        and linking vector embeddings to metadata chunk IDs.

        Args:
            file_path: The path to the document file.
        """
        if not os.path.exists(file_path):
            print(f"Error: File not found at {file_path}. Skipping indexing.")
            return

        print(f"Indexing document: {file_path}")

        # 1. Extract document-level metadata
        try:
            file_stat = os.stat(file_path)
            doc_metadata = {
                "file_path": file_path,
                "filename": os.path.basename(file_path),
                "file_type": os.path.splitext(file_path)[1].lower(),
                "size": file_stat.st_size,
                "creation_time": datetime.fromtimestamp(file_stat.st_ctime).isoformat(),
                "modification_time": datetime.fromtimestamp(file_stat.st_mtime).isoformat(),
                "title": None, # Will try to get from unstructured later
                "author": None # Will try to get from unstructured later
            }
        except Exception as e:
            print(f"Error extracting file metadata for {file_path}: {e}")
            doc_metadata = {"file_path": file_path, "filename": os.path.basename(file_path)} # Basic metadata


        # 2. Parse the document
        print("Parsing document...")
        elements = parse_document(file_path)
        if not elements:
            print(f"No content parsed from {file_path}. Skipping indexing.")
            return

        # Attempt to get title and author from unstructured metadata if available
        if elements and hasattr(elements[0], 'metadata'):
             if elements[0].metadata.title:
                 doc_metadata["title"] = elements[0].metadata.title
             if elements[0].metadata.author:
                 doc_metadata["author"] = elements[0].metadata.author


        # 3. Add/Update document metadata in the store
        print("Adding/Updating document metadata...")
        document_id = self.metadata_store.add_document_metadata(doc_metadata)
        if document_id is None:
            print(f"Failed to add/update document metadata for {file_path}. Skipping indexing.")
            return
        print(f"Document ID: {document_id}")

        # 4. Chunk the document elements
        print("Chunking document...")
        chunks_data = chunk_document_elements(elements, self.chunk_size, self.chunk_overlap)
        if not chunks_data:
            print(f"No chunks created from {file_path}. Skipping indexing.")
            return
        print(f"Created {len(chunks_data)} chunks.")

        # 5. Prepare chunk data for metadata store (add document_id and index)
        chunk_data_for_db = []
        for i, chunk in enumerate(chunks_data):
            chunk_data_for_db.append({
                "document_id": document_id,
                "chunk_index": i,
                "content": chunk.get("content", ""), # Include content
                "chunk_type": chunk.get("type", "text"),
                "start_char": None, # Placeholder
                "end_char": None,   # Placeholder
                "metadata": chunk.get("metadata", {})
            })

        # 6. Add chunk data to the metadata store
        print("Adding chunk data to metadata store...")
        # Note: This currently adds new chunks without removing old ones for updated documents.
        # A more complete versioning system would handle this.
        self.metadata_store.add_chunk_metadata(document_id, chunk_data_for_db)

        # 7. Retrieve the newly added chunks to get their generated IDs
        print("Retrieving chunk IDs from metadata store...")
        stored_chunks = self.metadata_store.get_chunks_by_document_id(document_id)
        if len(stored_chunks) != len(chunks_data):
             print(f"Warning: Mismatch between created chunks ({len(chunks_data)}) and stored chunks ({len(stored_chunks)}). Aborting embedding.")
             return

        # Create mapping from chunk_index to chunk_id and extract texts/ids
        chunk_texts = []
        external_ids_for_vector_store = [] # These will be the actual chunk IDs from the DB
        for stored_chunk in stored_chunks:
            chunk_texts.append(stored_chunk["content"])
            external_ids_for_vector_store.append(stored_chunk["id"]) # Use the DB primary key ID


        # 8. Generate embeddings for the chunks
        print("Generating embeddings...")
        embeddings = self.embedding_model.generate_embeddings(chunk_texts)
        print(f"Generated {len(embeddings)} embeddings.")

        # Ensure embeddings and external_ids match
        if len(embeddings) != len(external_ids_for_vector_store):
            print(f"Warning: Number of embeddings ({len(embeddings)}) does not match number of chunk IDs ({len(external_ids_for_vector_store)}).")
            min_count = min(len(embeddings), len(external_ids_for_vector_store))
            embeddings = embeddings[:min_count]
            external_ids_for_vector_store = external_ids_for_vector_store[:min_count]


        # 9. Add embeddings to the vector store using chunk IDs as external IDs
        print("Adding embeddings to vector store...")
        self.vector_store.add_embeddings(embeddings, external_ids_for_vector_store)
        print(f"Finished indexing {file_path}. Vector store size: {self.vector_store.get_size()}")


# Example Usage (Updated to reflect changes)
if __name__ == "__main__":
    # This example requires dummy files and initialized components
    dummy_docx_path = "data/documents/dummy_index_v2.docx"
    test_vector_db_dir = "data/vector_db_test_indexing_v2"
    test_metadata_db_path = "data/metadata_test_indexing_v2/metadata.db"

    os.makedirs("data/documents", exist_ok=True)
    os.makedirs(os.path.dirname(test_metadata_db_path), exist_ok=True)
    os.makedirs(test_vector_db_dir, exist_ok=True)

    # Create a dummy docx file if it doesn't exist
    if not os.path.exists(dummy_docx_path):
        try:
            from docx import Document
            doc = Document()
            doc.add_paragraph("This is the first paragraph of the dummy docx for indexing v2.")
            doc.add_paragraph("This is the second paragraph, testing metadata integration.")
            doc.save(dummy_docx_path)
            print(f"Created dummy docx file: {dummy_docx_path}")
        except ImportError:
            print("python-docx not installed. Cannot create dummy docx.")
        except Exception as e:
            print(f"Error creating dummy docx: {e}")

    try:
        # Initialize components
        print("\nInitializing components for indexing test...")
        embedding_model = EmbeddingModel()
        vector_store = VectorStore(index_path=test_vector_db_dir)
        metadata_store = MetadataStore(db_path=test_metadata_db_path)

        # Initialize DocumentIndexer
        document_indexer = DocumentIndexer(vector_store, embedding_model, metadata_store, chunk_size=100, chunk_overlap=20)

        # Index the dummy document
        if os.path.exists(dummy_docx_path):
            document_indexer.index_document(dummy_docx_path)

        # Example search after indexing (using the Retriever)
        print("\n--- Running Retrieval Test ---")
        from .retriever import Retriever # Import Retriever here
        retriever = Retriever(embedding_model, vector_store, metadata_store)
        # Build BM25 index (optional for this test, but good practice)
        retriever.build_bm25_index()

        query_text = "tell me about the dummy document v2"
        print(f"\nGenerating embedding for query: '{query_text}'")
        query_embedding = embedding_model.generate_embeddings([query_text])[0]

        print(f"\nSearching the vector store for top 1 similar chunk...")
        # Use the retriever to search
        search_results = retriever.retrieve(query_text, k=1, search_type="vector") # Use vector search for simplicity

        print("\nSearch Results:")
        if search_results:
            for result in search_results:
                print(f"  Chunk ID: {result['id']}, Distance: {result.get('distance', 'N/A'):.4f}")
                print(f"  Content: {result['content']}")
                print(f"  Metadata: {result['metadata']}")
        else:
            print("No search results.")


    except ImportError as e:
        print(f"Could not import necessary modules for indexing test: {e}. Skipping test.")
    except Exception as e:
        print(f"Error during indexing test: {e}")
    finally:
        # Clean up dummy files and directory
        if os.path.exists(dummy_docx_path):
            os.remove(dummy_docx_path)
        if os.path.exists(test_vector_db_dir):
            import shutil
            shutil.rmtree(test_vector_db_dir)
        if os.path.exists(test_metadata_db_path):
             os.remove(test_metadata_db_path)
        if os.path.exists(os.path.dirname(test_metadata_db_path)):
             try:
                 os.rmdir(os.path.dirname(test_metadata_db_path))
             except OSError:
                 pass
        print(f"\nCleaned up test files and directories.")