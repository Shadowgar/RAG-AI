import os
from typing import List, Dict, Any
from unstructured.partition.auto import partition

def parse_document(file_path: str) -> List[Dict[str, Any]]:
    """
    Parses a document using unstructured.partition and extracts elements
    with content and metadata.

    Args:
        file_path: The path to the document file.

    Returns:
        A list of dictionaries, where each dictionary represents an element
        extracted by unstructured, including its text content and metadata.
    """
    if not os.path.exists(file_path):
        print(f"Error: File not found at {file_path}")
        return []

    try:
        # Use unstructured.partition.auto to handle different file types
        elements = partition(filename=file_path)

        # Convert elements to a list of dictionaries
        parsed_data: List[Dict[str, Any]] = []
        for element in elements:
            # Elements have attributes like .text, .type, .metadata
            parsed_data.append({
                "content": str(element), # Use str(element) to get the text content
                "type": element.category if hasattr(element, 'category') else 'unknown', # Use category for type
                "metadata": element.metadata.to_dict() if hasattr(element, 'metadata') and element.metadata else {}
            })
        return parsed_data

    except Exception as e:
        print(f"Error parsing document {file_path} with unstructured: {e}")
        return []

if __name__ == "__main__":
    # Example usage (requires a test document)
    # Create a dummy file for testing
    dummy_docx_path = "data/documents/dummy.docx"
    try:
        from docx import Document
        doc = Document()
        doc.add_paragraph("This is a sample paragraph.")
        doc.add_paragraph("This is another paragraph.")
        doc.save(dummy_docx_path)

        print(f"Parsing dummy docx: {dummy_docx_path}")
        parsed_elements = parse_document(dummy_docx_path)
        for i, element in enumerate(parsed_elements):
            print(f"--- Element {i+1} ---")
            print(f"Type: {element.get('type')}")
            print(f"Content: {element.get('content')[:200]}...") # Print first 200 chars
            print(f"Metadata: {element.get('metadata')}")

    except ImportError:
        print("python-docx not installed. Skipping dummy docx test.")
    except Exception as e:
        print(f"Error during dummy docx test: {e}")
    finally:
        # Clean up dummy file
        if os.path.exists(dummy_docx_path):
            os.remove(dummy_docx_path)
            print(f"Cleaned up {dummy_docx_path}")