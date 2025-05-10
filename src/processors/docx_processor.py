import os
from typing import List, Dict, Any
from docx import Document as PythonDocxDocument # Alias for direct python-docx usage
from docxtpl import DocxTemplate # For template-based processing
from .base_processor import DocumentProcessor

class DocxProcessor(DocumentProcessor):
    """
    Document processor for Word (.docx) files.
    Uses DocxTemplate if replacements are provided, otherwise uses python-docx.
    Extracts text and basic structure while attempting to preserve
    information relevant for later formatting.
    """

    def supports(self, file_path: str) -> bool:
        """
        Checks if the processor supports the given file type.
        """
        return os.path.splitext(file_path)[1].lower() == ".docx"

    def process(self, file_path: str, replacements: Dict[str, str] = None) -> List[Dict[str, Any]]:
        """
        Processes a .docx file and extracts content and metadata.
        """
        if not self.supports(file_path):
            raise ValueError("File type not supported by DocxProcessor")

        chunks: List[Dict[str, Any]] = []
        processed_document_paragraphs = []

        try:
            if replacements:
                # Use DocxTemplate for processing with replacements
                doc_template = DocxTemplate(file_path)
                doc_template.render(replacements)

                # Manually update formatting after rendering based on replacement keys
                # This ensures that if a key like "bold" was used in replacements,
                # the run containing the replaced text becomes bold.
                for paragraph in doc_template.paragraphs:
                    for run in paragraph.runs:
                        # Check if this run's text is a *value* from the replacements
                        if run.text.strip() in replacements.values():
                            for rep_key, rep_value in replacements.items():
                                if run.text.strip() == rep_value:
                                    if rep_key == "bold":
                                        run.bold = True
                                    elif rep_key == "italic":
                                        run.italic = True
                                    # Add other formatting keys if needed (e.g., underline)
                processed_document_paragraphs = doc_template.paragraphs
            else:
                # No replacements, use python-docx directly to read the document as is
                plain_document = PythonDocxDocument(file_path)
                processed_document_paragraphs = plain_document.paragraphs

            # Extract chunks from the (potentially modified) document
            for paragraph in processed_document_paragraphs:
                for run in paragraph.runs:
                    metadata = {
                        "bold": run.bold,
                        "italic": run.italic,
                        "font_name": run.font.name if run.font else None,
                        "font_size": run.font.size.pt if run.font and run.font.size else None,
                        "source": file_path
                    }
                    # Debug print for run content and metadata
                    print(f"Run content: {run.text}, Metadata: {metadata}")

                    chunks.append({
                        "content": run.text,
                        "type": "paragraph", # Could be more specific if table/etc.
                        "metadata": metadata
                    })
            return chunks

        except Exception as e:
            print(f"Error processing {file_path}: {e}")
            return []